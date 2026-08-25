# -*- coding: utf-8 -*-
"""Build the PS 26082 engineering journal: what we did, what failed, what we learned."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "AREE_PS26082_Engineering_Journal.docx"

INK = RGBColor(0x14, 0x18, 0x20)
BODY = RGBColor(0x33, 0x3C, 0x4A)
MUTE = RGBColor(0x77, 0x81, 0x90)
BLUE = RGBColor(0x14, 0x47, 0x6E)
RED = RGBColor(0x9E, 0x2B, 0x25)
GREEN = RGBColor(0x1E, 0x6B, 0x44)
AMBER = RGBColor(0xB0, 0x6A, 0x0F)

SERIF, SANS, MONO = "Georgia", "Segoe UI", "Consolas"

doc = Document()


def _rf(r, font):
    rPr = r._element.get_or_add_rPr()
    rf = rPr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), font)


def R(p, text, size=10, bold=False, italic=False, color=BODY, font=SANS,
      caps=False, space=None):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.all_caps = caps
    _rf(r, font)
    if space is not None:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(space * 20)))
        r._element.get_or_add_rPr().append(sp)
    return r


def P(before=0, after=5, line=1.22, indent=None, keep=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after, pf.line_spacing = Pt(before), Pt(after), line
    pf.keep_with_next = keep
    if indent is not None:
        pf.left_indent = Cm(indent)
    return p


def rule(color="C8CFD8", size=6, before=0, after=8):
    p = P(before=before, after=after, line=1.0)
    bdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), color)
    bdr.append(b)
    p._p.get_or_add_pPr().append(bdr)


def h1(text, newpage=True):
    if newpage:
        doc.add_page_break()
    p = P(before=0, after=3, keep=True)
    R(p, text, size=18, bold=True, color=INK, font=SERIF)
    rule("14476E", 14, after=9)


def field(label, text, color=BODY, labelcolor=MUTE):
    p = P(after=2, line=1.2, indent=0.0, keep=True)
    R(p, label, size=7.4, font=MONO, color=labelcolor, caps=True, space=0.8)
    p = P(after=6, line=1.24, indent=0.45)
    if isinstance(text, str):
        R(p, text, size=9.5, color=color)
    else:
        for t, b in text:
            R(p, t, size=9.5, bold=b, color=INK if b else color)


def entry(n, title, status, did, failed, response, lesson):
    """One journal entry. Same five fields every time so it reads as a log."""
    p = P(before=13, after=3, keep=True)
    R(p, f"{n:02d}   ", size=13, bold=True, color=BLUE, font=MONO)
    R(p, title, size=13, bold=True, color=INK, font=SERIF)

    scolor = {"WORKED": GREEN, "FAILED": RED, "FIXED": AMBER,
              "FALSIFIED": RED, "SHIPPED": GREEN}.get(status, MUTE)
    p = P(after=5, line=1.0, indent=0.45)
    R(p, status, size=7.6, bold=True, font=MONO, color=scolor, caps=True, space=1.0)

    field("what we did", did)
    if failed:
        field("what went wrong", failed, color=BODY, labelcolor=RED)
    if response:
        field("what we did about it", response)
    field("lesson", lesson, color=INK, labelcolor=BLUE)
    rule("E4E8ED", 4, before=2, after=2)


# ------------------------------------------------------------------ setup
st = doc.styles["Normal"]
st.font.name = SANS
st.font.size = Pt(10)
st.font.color.rgb = BODY
st.element.rPr.rFonts.set(qn("w:eastAsia"), SANS)
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.line_spacing = 1.22

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin, sec.bottom_margin = Cm(1.9), Cm(1.7)
sec.left_margin, sec.right_margin = Cm(2.1), Cm(2.1)

f = sec.footer.paragraphs[0]
f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
R(f, "PS 26082 engineering journal  ·  Team Devengers  ·  page ",
  size=7, font=MONO, color=MUTE)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
f._p.append(fld)

# ------------------------------------------------------------------ cover
p = P(after=4, line=1.0)
R(p, "PROBLEM STATEMENT 26082  ·  ENGINEERING JOURNAL", size=8, font=MONO,
  color=BLUE, bold=True, caps=True, space=1.5)
p = P(after=12, line=1.0)
R(p, "Team Devengers  ·  Sri Shakthi IET, Coimbatore", size=8, font=MONO,
  color=MUTE, caps=True, space=1.3)

p = P(after=3, line=1.0)
R(p, "What we did, what broke,", size=25, bold=True, color=INK, font=SERIF)
p = P(after=8, line=1.0)
R(p, "and what it taught us.", size=25, bold=True, color=INK, font=SERIF)
rule("14476E", 16, before=4, after=10)

p = P(after=8, line=1.3)
R(p, "A running log of the build. Every entry records what was attempted, "
     "what went wrong, what we changed in response, and the lesson worth "
     "keeping. The failures are the useful part: several of them are the "
     "reason the final system is defensible, and at least one of them is the "
     "strongest thing we can put in front of a reviewer.", size=10)

p = P(before=6, after=3, keep=True)
R(p, "How to read this", size=7.6, font=MONO, color=MUTE, caps=True, space=0.9)
for label, colour, meaning in (
        ("WORKED", GREEN, "did what it was supposed to, first time"),
        ("FIXED", AMBER, "broke, was diagnosed, and now works"),
        ("FAILED", RED, "did not work and was abandoned"),
        ("FALSIFIED", RED, "a hypothesis we tested and disproved"),
        ("SHIPPED", GREEN, "in the running application")):
    p = P(after=2, line=1.15, indent=0.45)
    R(p, f"{label:<12}", size=8.4, bold=True, font=MONO, color=colour)
    R(p, meaning, size=9)

# ------------------------------------------------------------------ entries
h1("Phase 1 — Understanding what we already had")

entry(1, "Audit of the existing AREE codebase", "WORKED",
      "Cloned github.com/jeyamoorthi/urbanlive-ai and read all 4,533 lines of "
      "backend Python before writing anything new, rather than trusting the "
      "pitch deck's description of the system.",
      None,
      None,
      [("Two findings changed the plan. ", True),
       ("backend/ingestion/weather_stream.py was 0 bytes — the system ingested "
        "no meteorology at all — and the 'forecast' was numpy.polyfit degree 1 "
        "on recent AQI, projected 5 to 30 minutes. For a problem statement "
        "about weather-chemistry coupling over 72 hours, the gap was much "
        "larger than the architecture assumed. Read the code, not the deck.",
        False)])

entry(2, "Designing λ, the feedback loop gain", "WORKED",
      "Formulated the aerosol-radiation-PBL feedback as a product of three "
      "log-elasticities with closed-loop amplification 1/(1−λ), the same "
      "algebra used for climate feedback factors.",
      None,
      None,
      "A differentiator has to come out of the problem statement's own words, "
      "not be bolted on. The PS explicitly says these feedback loops are "
      "ignored; λ was an attempt to measure exactly the thing it named.")

entry(3, "Synthetic self-test before touching real data", "WORKED",
      "Built a generator with a prescribed λ and deliberately oversized "
      "diurnal cycles, then checked the estimator recovered it. It did: true "
      "0.1386, recovered 0.1513, 9.2% error, all three sign checks passing.",
      None,
      None,
      [("This passed, and it still misled us. ", True),
       ("A self-test can only falsify an estimator — it can never confirm "
        "that real data contains what you hope. It validated the ARITHMETIC, "
        "not the premise. Worth saying out loud to a reviewer: it is the "
        "difference between a team that tested its method and one that "
        "believed it.", False)])

h1("Phase 2 — Getting the data")

entry(4, "ERA5 from the Copernicus CDS", "FIXED",
      "Wrote a CDS API client for 23 monthly ERA5 requests covering five "
      "winters of NCR meteorology.",
      "No Copernicus key was available, and CDS requests queue for minutes to "
      "hours even with one. The single most important variable — boundary "
      "layer height — was behind that gate, and the whole project was blocked "
      "on it.",
      "Found that Open-Meteo's archive API re-serves ERA5 over plain REST "
      "with no key, no queue, and an arbitrary date range in one call — "
      "including boundary_layer_height and, crucially, terrestrial_radiation "
      "for the clearness index. Four and a half years of hourly data for a "
      "9-point NCR grid arrived in 90 seconds.",
      "Before accepting that a project is blocked on a credential, check "
      "whether the same data is re-served somewhere without one. The CDS "
      "client stayed in the repo as the canonical path; the REST path is what "
      "actually unblocked the work.")

entry(5, "OpenAQ sensor selection — three separate failures", "FIXED",
      "Pulled hourly ground PM2.5 for reference-grade NCR stations across "
      "2020–2025.",
      [("Three distinct bugs, each of which returned plausible-looking but "
        "wrong data rather than erroring. ", False),
       ("(a) ", True),
       ("A location's date range covers every sensor ever sited there, while "
        "individual sensor ids retire — selecting on location dates returned "
        "zero rows for recent windows. ", False),
       ("(b) ", True),
       ("Yearly request chunks forced deep pagination and OpenAQ answered "
        "with 500s and 408s, silently losing three stations. ", False),
       ("(c) ", True),
       ("Deduplicating by location kept the longest-overlap sensor, which is "
        "the OLD generation — this truncated the record at Oct 2022 and "
        "quietly removed both holdout Novembers.", False)],
      "Probed every sensor for its own coverage and cached the result; "
      "switched to monthly chunks so one request always fits inside a single "
      "1000-row page; stopped deduplicating across sensor generations and "
      "added the continuously-operating US Embassy monitor as the anchor that "
      "spans OpenAQ's Nov 2022 – Feb 2025 gap in the Indian feed.",
      [("Bug (c) is the dangerous kind. ", True),
       ("Nothing failed. The pipeline ran, produced a clean-looking dataset, "
        "and had silently deleted the evaluation period. Always print the "
        "coverage of what you actually loaded, not just its row count — the "
        "panel builder now reports stations per season every run.", False)])

entry(6, "CPCB via data.gov.in", "FIXED",
      "Wrote a connector for the authoritative CPCB CAAQMS feed to replace "
      "WAQI, which gives neither real station names nor a usable observation "
      "timestamp.",
      "The endpoint answered a single request in 2–4 seconds but returned "
      "HTTP 502 with a zero-length body after ~60 seconds under a burst — "
      "never a 429, so ordinary rate-limit handling did not catch it.",
      "Added backoff that treats a 5xx with an empty body as rate limiting.",
      "Read the failure mode, not the documentation. An API that rate-limits "
      "by timing out instead of by status code will defeat any retry logic "
      "written from the spec.")

h1("Phase 3 — The hypothesis meets the data")

entry(7, "λ on five winters of real data", "FALSIFIED",
      "Ran the full pipeline: 39,432 hours of meteorology, 29,953 hours of "
      "PM2.5, 161 labelled episodes of which 39 locked in.",
      [("λ does not separate locked-in from ventilated episodes. "
        "AUC 0.504 — chance.", True)],
      "Withdrew the claim rather than tuning until something crossed a "
      "threshold. Then spent the effort on diagnosing why instead.",
      "The validation gate existed precisely to produce this answer in week "
      "one. It cost one day. The alternative was five people building a "
      "presentation around a number that would have collapsed under the first "
      "question from an NCMRWF reviewer.")

entry(8, "Diagnosing the failure", "WORKED",
      "Tested three candidate explanations rather than assuming one: "
      "measurement error, spatial mismatch, and wrong timescale.",
      None,
      None,
      [("The raw physics is textbook — PM2.5 falls 163 to 79 µg/m³ as BLH "
        "rises 44 to 1102 m. What is unmeasurable is the residual after the "
        "diurnal cycle is removed: reverse regression bounds the reliability "
        "ratio of ERA5 BLH anomalies at ", False),
       ("0.016", True),
       (", roughly 98% noise. And e1 flips sign between specifications "
        "(−0.043 pooled, +0.061 dense-network, +0.084 daily). The elasticity "
        "is not small because the physics is weak; it is small because the "
        "instrument is.", False)])

entry(9, "Reformulating to avoid PBL height entirely", "FAILED",
      "If ERA5 BLH is the problem, remove it. Defined a ventilation index "
      "from observed PM2.5 alone: the pre-dawn to afternoon drop, which is "
      "the atmosphere ventilating, measured by a reference-grade monitor.",
      "Also fails to predict lock-in: AUC 0.554. It behaves correctly as a "
      "ventilation measure — +0.21 with daytime max BLH, +0.15 with wind — "
      "it just does not forecast which episodes persist.",
      "Stopped trying variants. Four predictors had now been tested on the "
      "same 161 episodes, and each further attempt inflates the chance of "
      "finding something spurious.",
      [("Knowing when to stop is part of the method. ", True),
       ("Two failures with different mechanisms pointing the same way is "
        "evidence about the problem, not an invitation to try a fifth index.",
        False)])

h1("Phase 4 — The finding that rescued the project")

entry(10, "Asking whether lock-in is predictable at all", "WORKED",
      "Instead of another predictor, asked a structural question: is lock-in "
      "determined by the state at onset, or by what happens afterwards? "
      "Compared identical predictors read before and after onset.",
      None,
      None,
      [("The answer was unambiguous. Ventilation BEFORE onset: AUC 0.514, no "
        "skill. AFTER onset: 0.736, skilful. ", False),
       ("Lock-in is not a diagnosable regime — it is a forecastable outcome.",
        True),
       (" Two failed hypotheses became one measured finding, and it points "
        "straight at what the problem statement actually asked for: coupled "
        "forecasting.", False)])

entry(11, "Testing whether forecast skill survives", "WORKED",
      "The 0.736 assumed perfect knowledge of the future. Used Open-Meteo's "
      "previous-runs archive — the value actually forecast for each hour one "
      "and two days earlier — to test how much survives with only a forecast.",
      "The archive carries no boundary layer height at any lead, so the "
      "strongest predictor could not be backtested. Only wind was available, "
      "and only from Jan 2024, giving 39 episodes with 9 locked-in.",
      "Ran the paired test on wind: analysis AUC 0.678, 1-day forecast 0.794, "
      "2-day 0.811, with forecast-vs-analysis agreement r = 0.75 and RMSE "
      "0.79 m/s.",
      [("Forecast error does not break the chain — but the forecast "
        "'beating' perfect knowledge is not a real effect. With 9 locked-in "
        "episodes the AUC standard error is about 0.10. ", False),
       ("Our own script printed '166% of the ceiling'; that number is an "
        "artefact and must never appear on a slide.", True)])

h1("Phase 5 — Building the application")

entry(12, "Calibrating the shipped decision threshold", "SHIPPED",
      "Derived the operating point from the same episodes and the same "
      "holdout as every other result, and emitted the whole ROC curve plus "
      "three named points instead of one.",
      None,
      None,
      "The exchange rate between a missed episode and a false GRAP invocation "
      "is a policy judgement for CAQM, not a modelling choice. The system "
      "ships a default of 466 m²/s (hit 0.61, false alarm 0.19) and lets the "
      "operator select a different point — and states the caveat that the "
      "sample is small.")

entry(13, "The forecast-to-escalation chain", "SHIPPED",
      "Three modules: a 72-hour ventilation forecast with sustained-collapse "
      "detection, a decision layer that opens an escalation case, and four "
      "API endpoints. Verified live: state IMMINENT, collapse in 11.4 hours, "
      "intervention window 11.4 hours, case AWAITING_APPROVAL.",
      None,
      None,
      [("The trigger is a CONJUNCTION — observed PM2.5 above the episode "
        "threshold AND a forecast ventilation collapse — because the research "
        "showed neither is sufficient alone. ", False),
       ("Verified by the negative case: PM2.5 = 40 produces no case at all. "
        "A trigger you have not tested in its OFF state is not a trigger, "
        "it is an alarm.", True)])

entry(14, "Keeping the forecast layer free of Pathway", "SHIPPED",
      "Pathway ships Linux/macOS wheels only, so every existing route "
      "degrades to engine_unavailable on Windows. The new routes were built "
      "without that dependency.",
      None,
      None,
      "The forecast layer is upstream of the streaming layer in the data "
      "flow, so it should not be downstream of it in the dependency graph. "
      "That reasoning — rather than convenience — is why the PS 26082 "
      "deliverable runs on the machines the team develops on.")

h1("Phase 6 — Making it real, and what that exposed")

entry(15, "Wiring live ground observations into the escalation", "FIXED",
      "The forecast half of the trigger was real, but the observed half was a "
      "hand-typed query parameter. Added a module reading the live CPCB/DPCC "
      "network so the conjunction runs on measured air.",
      "The first implementation queried OpenAQ's /parameters/2/latest with a "
      "bbox for the NCR domain and reported 198 stations. That endpoint "
      "SILENTLY IGNORES both bbox and coordinates+radius: the rows came from "
      "South Korea, Lithuania and China. The “Delhi composite” was a "
      "global median.",
      "Switched to /v3/locations?bbox= which does honour the box, then read "
      "each active location individually. Every returned coordinate is now "
      "re-checked against the NCR box as a hard filter.",
      [("It looked entirely plausible — a sensible station count and a "
        "sensible concentration — which is exactly what made it "
        "dangerous. ", False),
       ("An API that ignores a filter rather than rejecting it hands you a "
        "confident wrong answer. Verify the data you got back, not the query "
        "you sent.", True),
       (" It was only caught because someone asked to see the individual "
        "stations.", False)])

entry(16, "Two station counters disagreeing on one screen", "FIXED",
      "The page showed a live forecast and an observation panel reporting 33 "
      "stations.",
      "The global status strip sat directly above it reporting 0 / 0 stations "
      "and ENGINE NOT LOADED, because it reads the Pathway engine. Two "
      "contradictory numbers on one screen read as fabrication.",
      "Made the header and status strip report the subsystem the CURRENT view "
      "actually depends on, and added a panel naming both data sources "
      "explicitly: the forecast uses a weather model and no stations at all; "
      "the observation side uses the ground network.",
      "Chrome that reports a different subsystem than the page is worse than "
      "no chrome. If a number on screen contradicts another number on the same "
      "screen, users are right to distrust both — and they should.")

entry(17, "Running Pathway under WSL", "FAILED",
      "Attempted to start the Ubuntu-22.04 distro so the streaming engine "
      "could run and restore the other three tabs.",
      "The distro will not mount: E_ACCESSDENIED on its vhdx. Isolated to that "
      "disk specifically — the docker-desktop distro launches fine. The "
      "vhdx sits at D:/WSL/Ubuntu after being moved off the default "
      "location, and its ACL is missing the generic NT VIRTUAL MACHINE/"
      "Virtual Machines entry that WSL needs after a move. Fixing it requires "
      "an elevated shell, which the working session did not have.",
      "Stopped rather than attempting to escalate privileges, and handed over "
      "the exact icacls command. Then solved the underlying problem a "
      "different way — see the next entry.",
      "A blocker that needs someone else's authority is a handover, not a "
      "puzzle to force. But being blocked on the stated approach is not the "
      "same as being blocked on the goal.")

entry(18, "Direct-mode engine: all four tabs without Pathway", "SHIPPED",
      "Rather than wait on WSL, checked which modules actually need the "
      "Pathway runtime. Only rag/advisory_engine.py imports it. The GRAP state "
      "machine, persistence tracking, causal attribution, satellite fire "
      "intelligence and the report generator are all pure Python.",
      None,
      "Built fallback_engine.py: samples the live CPCB network and drives the "
      "SAME state machine, exposing the same module attributes the API layer "
      "already reads. engine.load_engine() falls back to it when the Pathway "
      "import fails. Every state carries mode=“direct”, and the "
      "subsystems it genuinely cannot provide (policy RAG, event-time windows, "
      "carbon tracking) are reported as unavailable rather than faked. Result: "
      "52 stations live, real GRAP stages, all four tabs working.",
      [("The dependency was never as deep as the failure made it look. ",
        True),
       ("Three tabs were dark because of one import in one module, and "
        "displaying a station's air quality never needed a streaming runtime "
        "— it needed observations and a state machine. Before accepting "
        "that a platform constraint blocks a feature, check how much of the "
        "feature actually touches the constraint.", False)])

# ------------------------------------------------------------------ closing
h1("What we would tell another team")

for text in [
    [("Build the gate before the pitch. ", True),
     ("A one-week experiment that can kill the central idea is worth more "
      "than a month of building on it. Ours killed it, and the project is "
      "stronger for it.", False)],
    [("A passing self-test is not evidence about the world. ", True),
     ("It tells you the code computes what you wrote down. Whether reality "
      "contains that quantity is a separate question and needs separate data.",
      False)],
    [("Print what you loaded, not just how much. ", True),
     ("The bug that silently deleted both holdout months did not raise an "
      "error. It produced a clean dataset of the wrong period.", False)],
    [("Negative results are presentable. ", True),
     ("'We measured whether lock-in can be diagnosed from the present state; "
      "it cannot' is a stronger sentence than any hypothesis we could have "
      "asserted, because it is the one a reviewer cannot knock down.", False)],
    [("Test the trigger in its off state. ", True),
     ("Most demos only show the alarm firing.", False)],
    [("Verify the data you got back, not the query you sent. ", True),
     ("An API that ignores a filter instead of rejecting it will hand you a "
      "plausible, confident, wrong answer.", False)],
    [("Check how much of the feature actually touches the blocker. ", True),
     ("Three tabs were dark because of one import in one module.", False)],
]:
    p = P(after=7, line=1.28, indent=0.45)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    R(p, "—   ", size=9.5, font=MONO, color=BLUE)
    for t, b in text:
        R(p, t, size=9.8, bold=b, color=INK if b else BODY)

P(after=8)
rule("C8CFD8", 4, before=8, after=10)
p = P(after=0, line=1.3)
R(p, "Every number in this journal is reproducible from the scripts in "
     "research/ps26082 on branch ps26082-coupled-forecast.",
  size=10, italic=True, color=INK, font=SERIF)
p = P(before=8, after=0)
R(p, "Team Devengers   ·   PS 26082   ·   Journal revision A",
  size=7.5, font=MONO, color=MUTE, caps=True, space=0.6)

doc.save(OUT)
print("saved:", OUT)
