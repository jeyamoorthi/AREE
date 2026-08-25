# -*- coding: utf-8 -*-
"""Build the PS 26082 architecture document (.docx, then Word exports the PDF)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(r"d:\AZUO\chat\aree-ps26082\docs\AREE_PS26082_Architecture.docx")

INK = RGBColor(0x14, 0x18, 0x20)
BODY = RGBColor(0x33, 0x3C, 0x4A)
MUTE = RGBColor(0x77, 0x81, 0x90)
BLUE = RGBColor(0x14, 0x47, 0x6E)
RED = RGBColor(0x9E, 0x2B, 0x25)
GREEN = RGBColor(0x1E, 0x6B, 0x44)

SERIF = "Georgia"
SANS = "Segoe UI"
MONO = "Consolas"

doc = Document()


# ---------------------------------------------------------------- primitives

def _rfonts(run, font):
    rPr = run._element.get_or_add_rPr()
    rf = rPr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), font)


def R(p, text, size=10, bold=False, italic=False, color=BODY, font=SANS,
      caps=False, space=None):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.all_caps = caps
    _rfonts(r, font)
    if space is not None:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(space * 20)))
        r._element.get_or_add_rPr().append(sp)
    return r


def P(before=0, after=6, line=1.25, indent=None, align=None, keep=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after, pf.line_spacing = Pt(before), Pt(after), line
    pf.keep_with_next = keep
    if indent is not None:
        pf.left_indent = Cm(indent)
    if align is not None:
        p.alignment = align
    return p


def para(text, **kw):
    """Plain body paragraph. Accepts a list of (text, bold) tuples for inline bold."""
    p = P(**{k: v for k, v in kw.items() if k in
             ("before", "after", "line", "indent", "align", "keep")})
    if isinstance(text, str):
        R(p, text)
    else:
        for item in text:
            if isinstance(item, tuple):
                R(p, item[0], bold=item[1],
                  color=INK if item[1] else BODY,
                  font=item[2] if len(item) > 2 else SANS)
            else:
                R(p, item)
    return p


def h1(text, num=None):
    doc.add_page_break()
    p = P(before=0, after=3, keep=True)
    if num:
        R(p, f"{num}   ", size=20, bold=True, color=BLUE, font=SERIF)
    R(p, text, size=20, bold=True, color=INK, font=SERIF)
    rule(color="14476E", size=14, after=10)


def h2(text):
    p = P(before=14, after=4, keep=True)
    R(p, text, size=12.5, bold=True, color=INK, font=SERIF)


def h3(text):
    p = P(before=10, after=3, keep=True)
    R(p, text, size=9.5, bold=True, color=BLUE, font=MONO, caps=True, space=0.8)


def rule(color="C8CFD8", size=6, before=0, after=8):
    p = P(before=before, after=after, line=1.0)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), color)
    bdr.append(b)
    pPr.append(bdr)


def bullets(items, indent=0.55):
    for it in items:
        p = P(after=3, line=1.2, indent=indent)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        R(p, "\u2022   ", size=9, color=BLUE, font=MONO)
        if isinstance(it, str):
            R(p, it, size=9.5)
        else:
            for item in it:
                if isinstance(item, tuple):
                    R(p, item[0], size=9.5, bold=item[1],
                      color=INK if item[1] else BODY,
                      font=item[2] if len(item) > 2 else SANS)
                else:
                    R(p, item, size=9.5)


def code(text, size=8.2):
    """Monospace block for diagrams, formulas and shell output."""
    for line in text.strip("\n").split("\n"):
        p = P(after=0, line=1.02, indent=0.35)
        R(p, line if line else " ", size=size, font=MONO, color=INK)
    P(after=6, line=1.0)


def callout(title, text, color=RED):
    p = P(before=8, after=2, line=1.2, indent=0.3, keep=True)
    R(p, title, size=9, bold=True, font=MONO, caps=True, color=color, space=0.7)
    p = P(after=8, line=1.22, indent=0.3)
    if isinstance(text, str):
        R(p, text, size=9.5)
    else:
        for item in text:
            if isinstance(item, tuple):
                R(p, item[0], size=9.5, bold=item[1], color=INK if item[1] else BODY)
            else:
                R(p, item, size=9.5)


def set_cell_borders(cell, **kw):
    tcPr = cell._tc.get_or_add_tcPr()
    bs = tcPr.find(qn("w:tcBorders"))
    if bs is None:
        bs = OxmlElement("w:tcBorders")
        tcPr.append(bs)
    for edge in ("top", "left", "bottom", "right"):
        spec = kw.get(edge)
        if spec is None:
            continue
        tag = f"w:{edge}"
        el = bs.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            bs.append(el)
        sz, col = spec
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), col)


def table(headers, rows, widths_cm, font_size=8.5, colors=None):
    """Borderless table with hairline row separators."""
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tblPr = t._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    nb = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        nb.append(el)
    tblPr.append(nb)
    mar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", 60), ("left", 0), ("bottom", 60), ("right", 130)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)
    t.autofit = False
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)
    for row in t.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)

    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_borders(c, bottom=(10, "9AA5B2"))
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        R(p, htxt, size=7.6, font=MONO, color=MUTE, caps=True, space=0.7)

    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            if ri > 1:
                set_cell_borders(c, top=(4, "DDE2E9"))
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.12
            col = BODY
            if colors and colors.get(ci):
                col = colors[ci]
            R(p, str(val), size=font_size, color=col,
              bold=(ci == 0), font=SANS)
    P(after=4)
    return t


# ---------------------------------------------------------------- page setup

st = doc.styles["Normal"]
st.font.name = SANS
st.font.size = Pt(10)
st.font.color.rgb = BODY
st.element.rPr.rFonts.set(qn("w:eastAsia"), SANS)
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.line_spacing = 1.25

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(1.8)
sec.left_margin, sec.right_margin = Cm(2.2), Cm(2.2)

# footer with page number
ftr = sec.footer.paragraphs[0]
ftr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
R(ftr, "AREE / PS 26082 architecture  \u00b7  Team Devengers  \u00b7  page ",
  size=7, font=MONO, color=MUTE)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
ftr._p.append(fld)


# ================================================================= COVER

p = P(after=4, line=1.0)
R(p, "PROBLEM STATEMENT 26082", size=8, font=MONO, color=BLUE, bold=True,
  caps=True, space=1.6)
p = P(after=14, line=1.0)
R(p, "Ministry of Earth Sciences  \u00b7  NCMRWF  \u00b7  Disaster Management  \u00b7  Software",
  size=8, font=MONO, color=MUTE, caps=True, space=1.4)

p = P(after=2, line=0.95)
R(p, "AREE 2.0", size=34, bold=True, color=INK, font=SERIF)
p = P(after=3, line=1.15)
R(p, "Air Pollution\u2013Weather Coupled Forecasting System for Delhi NCR",
  size=13, italic=True, color=BODY, font=SERIF)
rule(color="14476E", size=18, before=6, after=8)

p = P(after=16, line=1.3)
R(p, "Technical architecture, module specification, algorithm selection and "
     "team allocation. Includes the validated ", size=10)
R(p, "aerosol\u2013boundary-layer feedback loop gain (\u03bb)", size=10, bold=True, color=INK)
R(p, " that is the scientific core of this submission, its evidence base, and "
     "the runnable pipeline that proves or kills it before any model is built.", size=10)

h3("What this document decides")
bullets([
    [("The scientific differentiator", True), " \u2014 \u03bb, the loop gain of the "
     "aerosol\u2013radiation\u2013PBL feedback, and the moment it crosses into "
     "self-amplification."],
    [("What is reused from the existing AREE codebase", True),
     " \u2014 audited file by file, not assumed."],
    [("The forecasting algorithm", True), " \u2014 and why every alternative was "
     "rejected."],
    [("The module boundaries", True), " \u2014 five owners, minimal coupling, "
     "each independently testable."],
    [("The evaluation protocol", True), " \u2014 the numbers NCMRWF will ask for."],
])

P(after=10)
callout("status at time of writing",
        [("The pipeline in "), ("aree-ps26082/scripts/", False, MONO),
         (" is written and passes its synthetic self-test: the estimator recovers "
          "a prescribed \u03bb of 0.1386 as 0.1513, a 9.2% error, with all three "
          "physical sign checks passing. It has "), ("not", True),
         (" yet been run on real ERA5 data \u2014 that requires a Copernicus CDS "
          "key and is step one of the work plan in section 11.")],
        color=GREEN)

# ================================================================= 1

h1("The correction this architecture makes", "1")

para([("PS 26082 is owned by NCMRWF \u2014 India's operational numerical weather "
       "prediction centre \u2014 and sits under the "), ("Disaster Management", True),
      (" theme. Both facts change what the deliverable is.")])

para("The previous AREE positioning was built for SIH25216, a policy-dashboard "
     "problem where the winning wedge was governance: everyone else stopped at a "
     "recommendation, so AREE owned the regulatory execution layer after it. That "
     "wedge does not transfer. Here the evaluators are atmospheric scientists and "
     "the deliverable they are judging is forecast skill under coupling. Leading "
     "with escalation workflow would read as last year's project brought to a "
     "different problem.")

h2("The inversion")
para("AREE does not die. It changes position in the stack. It stops being the "
     "product and becomes the operational layer of a forecasting system \u2014 what "
     "converts a physical prediction into lead time someone can act on, which is "
     "precisely what Disaster Management means.")

code("""
   OLD  (SIH25216)                    NEW  (PS 26082)

   AQ observation                     physical atmosphere
        |                                    |
        v                                    v
   +-----------+                     +------------------+
   |   AREE    |  <-- the product    | COUPLED FORECAST |  <-- the product
   | DETECT    |                     | weather <-> aerosol
   | DECIDE    |                     +--------+---------+
   | ACT       |                              |
   +-----------+                              v
        |                             +------------------+
        v                             | REGIME DIAGNOSTIC|
   regulatory action                  | loop gain lambda |
                                      | lock-in time     |
                                      | uncertainty      |
                                      +--------+---------+
                                               |
                                               v
                                      intervention window
                                               |
                                               v
                                      +------------------+
                                      |      AREE        |  <-- the delivery
                                      | rules, authority |      mechanism
                                      | escalation, audit|
                                      +------------------+
""")

h2("The one-line positioning")
p = P(before=6, after=8, indent=0.3, line=1.3)
R(p, "Everyone forecasts the pollution level. We forecast the moment the "
     "atmosphere stops responding to intervention.",
  size=12.5, italic=True, color=BLUE, font=SERIF)

callout("do not use this line yet",
        "It is only true if section 4's validation returns a usable separation on "
        "real data. If it does not, the honest fallback is a coupled forecasting "
        "system with an explicit feedback diagnostic and no lock-in claim. Find "
        "that out privately, in week one, not in front of judges.")

# ================================================================= 2

h1("Audit of the existing AREE codebase", "2")

para([("Findings from "), ("github.com/jeyamoorthi/urbanlive-ai", False, MONO),
      (" at commit "), ("e876f6d", False, MONO), (", cloned and read in full. "
      "This section exists because the reuse plan has to be based on what the "
      "code does, not on what the pitch deck says it does.")])

h2("What is genuinely there and worth keeping")

table(
    ["Component", "File", "Assessment"],
    [
        ["Pathway streaming DAG", "backend/app.py",
         "Real. pw.io.python.read with custom ConnectorSubject, genuine "
         "pw.temporal.sliding windows at 3/5/15 min with 1 min hop, "
         "pw.reducers aggregation. Not a timer loop dressed as streaming."],
        ["GRAP state machine", "streaming/state_machine.py",
         "Clean deterministic implementation. Hysteresis requires "
         "HYSTERESIS_CONFIRMATIONS consecutive agreeing windows before a stage "
         "transition. Per-station state. Directly reusable."],
        ["Persistence tracker", "streaming/state_machine.py",
         "Counts consecutive high-AQI windows before triggering. Correct "
         "pattern, wrong timescale for this PS \u2014 see below."],
        ["Causal attribution", "streaming/risk_engine.py",
         "CausalAttributionEngine + TransportVectorModel: FIRMS fire counts "
         "joined with wind bearing. Reusable as an emission-proxy input."],
        ["FIRMS ingestion", "ingestion/firms_stream.py",
         "Working NRT satellite fire pull with confidence filtering and bbox "
         "query. Keep as-is."],
        ["Regulatory RAG", "rag/advisory_engine.py",
         "Document-grounded advisory generation over the policies directory. "
         "Keep, demote in prominence."],
        ["Audit and reporting", "report_generator.py",
         "PDF brief generation. Keep as the order-export path."],
        ["Next.js operations UI", "frontend/",
         "Next 16, React 19, Recharts, Leaflet. Real dashboard. The spatial "
         "forecast layer plugs into the existing Leaflet map."],
    ],
    [3.2, 3.6, 9.8],
)

h2("What is missing, and it is more than the pitch assumes")

callout("critical finding",
        [("backend/ingestion/weather_stream.py", False, MONO),
         (" is "), ("zero bytes", True), (". The current system ingests no "
          "meteorology at all. Temperature, humidity, pressure, boundary layer "
          "height and radiation are absent from the codebase entirely. For a "
          "problem statement whose entire subject is weather\u2013chemistry "
          "coupling, this is the gap that defines the work.")])

table(
    ["Gap", "Current state", "Required by PS 26082"],
    [
        ["Meteorology", "weather_stream.py is empty (0 lines). Only wind speed "
         "and bearing arrive, embedded in the WAQI payload.",
         "Full met state: T, RH, wind, pressure, PBL height, shortwave down, "
         "clear-sky shortwave."],
        ["Forecast method",
         "compute_short_term_forecast() in app.py: numpy.polyfit degree 1 on "
         "recent AQI, extrapolated linearly.",
         "Coupled physical evolution over 72 hours."],
        ["Forecast horizon", "5 and 30 minutes.",
         "72 hours \u2014 a 144x extension. Linear extrapolation is not a "
         "candidate at this horizon."],
        ["Spatial structure", "5 point stations, only 2 inside Delhi NCR.",
         "Gridded NCR domain."],
        ["Coupling", "None. Chemistry and meteorology never interact.",
         "Two-way aerosol\u2013radiation\u2013PBL feedback \u2014 the explicit "
         "subject of the PS."],
        ["Persistence timescale",
         "PERSISTENCE_THRESHOLD = 3 windows of 3 minutes.",
         "Meteorological episodes persist over 12\u201372 hours. The mechanism "
         "is right, the constant is three orders of magnitude too small."],
        ["Verification", "None in the repository.",
         "MAE, RMSE, bias against persistence and an external baseline."],
    ],
    [2.9, 6.4, 7.3],
)

h2("Two claims to stop making immediately")
bullets([
    [("The 23 ms figure.", True), " That is React/API round-trip latency. It has "
     "nothing to do with forecasting. Quoted anywhere near a 72-hour physical "
     "forecast, an NWP scientist reads it as a category error and the room is "
     "lost. Report it, if at all, as \u201cdecision-layer response latency\u201d "
     "in a separate section."],
    [("\u201cStreams every thirty seconds.\u201d", True), " AQI_POLL_INTERVAL = 30 "
     "is the poll interval. WAQI republishes roughly hourly, which the codebase "
     "already acknowledges \u2014 FRESH_DATA_THRESHOLD_SECONDS is set to 90 "
     "minutes. Polling frequency is not measurement freshness. Report observed_at, "
     "received_at and processed_at separately and display data age. That is more "
     "credible, not less."],
])

# ================================================================= 3

h1("The scientific core: the feedback loop gain \u03bb", "3")

h2("3.1  The mechanism, in the problem statement's own words")

para([("The PS says: "), ("\u201cdense concentrations of aerosols (PM2.5) block "
      "sunlight, altering local temperatures, wind patterns, and planetary "
      "boundary layer (PBL) heights\u2026 Ignoring these coupled "
      "meteorological-chemical feedback loops leads to significant "
      "inaccuracies.\u201d", False, SERIF)])

para("That sentence describes a positive feedback. Written out:")

code("""
        PM2.5 rises
             |
             v
   surface shortwave falls        (aerosol scattering + absorption)
             |
             v
   sensible heat flux falls       (less energy into the surface layer)
             |
             v
   turbulent kinetic energy falls (weaker buoyancy production)
             |
             v
   PBL height collapses           (shallower mixed layer)
             |
             v
   mixing volume shrinks
             |
             +---------------> PM2.5 rises further   [loop closes]
""")

para("Every other team will treat this coupling as an implementation detail on "
     "the way to a better AQI number. Their output is a line graph. The "
     "differentiator proposed here is to treat the loop as a diagnosable "
     "state with a critical threshold, and to report the threshold crossing "
     "as the primary product.")

h2("3.2  Formulation")

para("Take the three log-sensitivities (elasticities) around the ring:")

code("""
    e1 = d ln C / d ln H      PM2.5 response to boundary-layer depth
    e2 = d ln H / d ln S      PBL response to surface shortwave
    e3 = d ln S / d ln C      shortwave response to aerosol load

    lambda = e1 * e2 * e3
""")

para("Expected signs follow from elementary physics:")

table(
    ["Term", "Sign", "Physical basis", "Expected magnitude"],
    [
        ["e1", "negative", "Box dilution. For fixed emission E, C ~ E / (H\u00b7U), "
         "so d ln C / d ln H = \u22121 exactly.", "\u2248 \u22121"],
        ["e2", "positive", "Encroachment growth. Mixed-layer depth scales with the "
         "time-integrated surface heat flux, H ~ S^\u00bd.", "\u2248 +0.5"],
        ["e3", "negative", "Beer\u2013Lambert extinction. Optical depth rises with "
         "column aerosol mass, transmitted flux falls exponentially.",
         "\u22120.1 to \u22120.4"],
    ],
    [1.6, 2.0, 8.2, 4.8],
)

para([("The product of ("), ("\u2212", True), (")("), ("+", True), (")("),
      ("\u2212", True), (") is "), ("positive", True), (". A positive loop gain "
      "is a self-reinforcing feedback. The closed-loop amplification of any "
      "perturbation is then")])

code("""
    A = 1 / (1 - lambda)                for lambda < 1
""")

para([("which is the standard linear feedback-factor form used throughout "
       "geophysics \u2014 the same algebra as the climate feedback parameter "
       "(Roe, "), ("Feedbacks, Timescales, and Seeing Red", True),
      (", Annual Review of Earth and Planetary Sciences 37:93\u2013115, 2009). "
       "As \u03bb approaches 1 the amplification diverges: the atmosphere stops "
       "damping perturbations and the episode locks in.")])

callout("why elasticities and not raw partial derivatives",
        "Elasticities are dimensionless, so \u03bb is dimensionless, comparable "
        "across cities, seasons and years, and independent of the units anyone "
        "reports PM2.5 or PBL height in. A product of raw partial derivatives "
        "would carry units of (\u03bcg m\u207b\u00b3 / m)(m / W m\u207b\u00b2)"
        "(W m\u207b\u00b2 / \u03bcg m\u207b\u00b3), which happens to cancel \u2014 "
        "but only by construction, and the intermediate quantities would not be "
        "interpretable on their own.", color=BLUE)

h2("3.3  Evidence base")

para("The mechanism is established in the peer-reviewed literature. What is "
     "novel here is operationalising it as a live, forecastable scalar with a "
     "critical threshold \u2014 not the physics itself. That distinction must be "
     "stated explicitly in the presentation; claiming \u03bb as a published "
     "standard index would not survive a single question.")

table(
    ["Claim in the architecture", "Supporting literature", "What it constrains"],
    [
        ["Aerosols measurably dim the surface over Delhi / the Indo-Gangetic "
         "Plain during haze episodes.",
         "WRF-Chem studies of aerosol optical properties and surface shortwave "
         "over Delhi report reductions of order 25\u201380 W m\u207b\u00b2 during "
         "polluted episodes.",
         "Magnitude of e3."],
        ["Reduced surface heating suppresses turbulence and collapses the mixed "
         "layer.",
         "Winter Fog Experiment (WiFEX) and IGP haze studies report daytime "
         "maximum PBL height falling from roughly 1.3 km in clean conditions "
         "toward roughly 0.6 km under heavy aerosol loading.",
         "Magnitude of e2 and the operating range of H."],
        ["Absorbing aerosol aloft stabilises the column \u2014 the \u201cdome "
         "effect\u201d.",
         "Ding et al., Geophysical Research Letters, 2016, on black-carbon "
         "induced boundary-layer suppression; Peta\u0308ja\u0308 et al., "
         "Scientific Reports, 2016, on aerosol\u2013boundary-layer feedback in "
         "polluted megacities.",
         "Existence and sign of the loop."],
        ["Compressed, humid boundary layers accelerate secondary aerosol "
         "formation, strengthening the loop.",
         "Studies of the secondary-aerosol positive feedback during severe haze "
         "over the Beijing\u2013Tianjin\u2013Hebei region; Delhi speciation work "
         "showing organics at roughly 50\u201369% of non-refractory submicron "
         "mass with a large chloride fraction from solid-fuel combustion.",
         "Why e3 must not be estimated with RH controlled out."],
        ["Feedback factors compose as A = 1/(1\u2212\u03bb).",
         "Roe, Annual Review of Earth and Planetary Sciences, 2009 \u2014 the "
         "canonical treatment of feedback composition in geophysical systems.",
         "The algebra of the threshold claim."],
        ["Ventilation coefficient (PBLH \u00d7 wind) is the conventional "
         "dispersion metric.",
         "Standard in operational air-quality meteorology.",
         "The baseline \u03bb must outperform."],
    ],
    [4.3, 7.0, 5.3],
)

callout("verify before submission",
        "Every citation above should be pulled, read and page-checked by the "
        "team before it appears on a slide. They are given here as the correct "
        "places to look, from a literature survey \u2014 not as verified page "
        "references. A wrong citation in front of NCMRWF costs more than no "
        "citation.")

h2("3.4  The trap the estimator is built to survive")

para("PM2.5, PBL height and surface shortwave all have large, in-phase diurnal "
     "cycles. A naive regression of ln C on ln H recovers the diurnal cycle, not "
     "the feedback \u2014 and returns a beautiful, meaningless e1 near \u22121 "
     "every single time. This is the single most likely way for this project to "
     "produce a confident wrong answer.")

para([("The fix, applied identically to all four variables in "),
      ("log_anomaly()", False, MONO), (": every variable is converted to a log "
      "anomaly from its own (month, hour-of-day) climatology before any "
      "regression. What remains is the perturbation about normal diurnal "
      "evolution \u2014 exactly the quantity a feedback gain is defined on.")])

h2("3.5  Self-test result")

para("The estimator was tested against synthetic data with a prescribed loop "
     "gain, generated with deliberately large diurnal cycles so that a "
     "de-seasonalisation failure would show up as a test failure rather than "
     "passing quietly.")

code("""
    synthetic ground truth:  e1=-0.90   e2=+0.55   e3=-0.28
                             TRUE lambda = 0.1386

    physical sign checks
    --------------------------------------------------------------
    PASS  e1 = dlnC/dlnH  should be NEGATIVE (dilution)      -1.961
    PASS  e2 = dlnH/dlnS  should be POSITIVE (radiative)     +0.687
    PASS  e3 = dlnS/dlnC  should be NEGATIVE (attenuation)   -0.160

    recovered lambda  +0.1513   (median of rolling estimate)
    relative error    9.2%
    PASS - estimator recovers the prescribed loop gain

    separation on synthetic episodes
      lambda                          AUC 0.733   median lead time 12 h
      negative ventilation coeff.     AUC 0.433   median lead time 22 h
""")

callout("known limitation \u2014 report this, do not hide it",
        [("The individual elasticities are biased: e1 came back at \u22121.96 "
          "against a true \u22120.90. This is "), ("simultaneity bias", True),
         (", the classic endogeneity problem when ordinary least squares is "
          "applied to a closed loop \u2014 each regressor is correlated with the "
          "error term because the system is solved jointly. The product \u03bb "
          "survives because the biases partially cancel, but this is luck, not "
          "design. The correct fix is instrumental-variable or "
          "simultaneous-equations estimation, using an instrument that shifts one "
          "variable without acting through the others \u2014 synoptic wind is the "
          "natural candidate. This is a real piece of work and it belongs to "
          "Module B.")])

# ================================================================= 4

h1("Validation protocol: the go / no-go gate", "4")

para("The central claim is falsifiable and must be falsified or confirmed before "
     "anything is built on it. The pipeline does this without a forecast model, "
     "using observations alone, which also means the result survives even if the "
     "emulator underdelivers.")

h2("The circularity that had to be designed out")

callout("why \u03bb is computed upstream, not downstream",
        "In the earlier architecture sketch \u03bb was computed from the "
        "emulator's forecast fields. That is circular: if the emulator predicts "
        "PM2.5 rising and PBL collapsing, \u03bb mechanically comes out high, and "
        "the diagnostic is just a re-reading of the model's own output. The first "
        "question a reviewer asks is whether \u03bb says anything the PM2.5 curve "
        "did not. Computing \u03bb from observations alone, over past winters, "
        "before any model exists, removes both the circularity and the single "
        "point of failure.")

h2("The test")

table(
    ["Step", "What happens", "Guard against"],
    [
        ["Label episodes",
         "Contiguous runs above a PM2.5 threshold, short gaps bridged. "
         "\u201cLocked in\u201d requires BOTH duration \u2265 48 h AND peak "
         "\u2265 250 \u03bcg m\u207b\u00b3.",
         "Labels never see \u03bb. A long mild event is not a lock-in; a brief "
         "sharp local spike is not either."],
        ["Read the predictor",
         "Maximum \u03bb in the 24 h BEFORE onset.",
         "Reading before onset is what makes it a forecast test rather than a "
         "description."],
        ["Separation",
         "ROC AUC via the Mann\u2013Whitney identity, hand-implemented.",
         "No sklearn dependency, no hidden weighting; the computation is "
         "auditable line by line."],
        ["Threshold sweep",
         "Full confusion table across all candidate thresholds; Youden J "
         "reported but not imposed.",
         "The operating point is a policy judgement about the cost of a false "
         "GRAP invocation, not a modelling choice."],
        ["Lead time",
         "Hours between first threshold crossing and onset.",
         "A diagnostic firing 2 h ahead is a description. One firing 18 h ahead "
         "is decision support."],
        ["Baseline",
         "The identical test using negative ventilation coefficient.",
         "If \u03bb cannot beat PBLH \u00d7 wind, it adds nothing and should be "
         "dropped."],
        ["Holdout",
         "November 2023 and November 2024 excluded from every fit.",
         "The only numbers quoted publicly come from episodes the fit never saw."],
    ],
    [2.6, 6.6, 7.4],
)

h2("Decision rule")

table(
    ["Held-out AUC", "Verdict", "Action"],
    [
        ["\u2265 0.75", "\u03bb separates the classes.",
         "Proceed. \u03bb becomes the headline and the emulator is built to "
         "forecast it."],
        ["0.60 \u2013 0.75", "Weak separation.",
         "Fix the estimator first \u2014 instrumental variables for the "
         "simultaneity bias, ceilometer-corrected PBLH. Do not present the "
         "lock-in claim yet."],
        ["< 0.60", "\u03bb does not separate the classes.",
         "Drop the lock-in claim entirely. Fall back to a coupled forecasting "
         "system with an explicit feedback diagnostic \u2014 still responsive to "
         "the PS, still ahead of an uncoupled ML forecaster."],
    ],
    [3.0, 4.2, 9.4],
    colors={1: INK},
)

# ================================================================= 5

h1("Algorithm selection", "5")

h2("5.1  The forecasting engine")

para("The PS names WRF-Chem, and its own wording gives the escape hatch: "
     "\u201cWRF-Chem or similar open-source coupled frameworks\u201d. That "
     "phrasing is deliberate and should be taken.")

h3("Rejected: WRF-Chem in the operational loop")
bullets([
    "Requires an emissions inventory, chemical boundary conditions, a chemical "
    "mechanism and hours of HPC time per forecast cycle.",
    "Cannot be demonstrated live. Every team that promises this will be showing "
    "pre-computed output, a toy domain, or nothing.",
    "A five-person student team cannot operate it on a hackathon timeline.",
])

h3("Rejected: pure sequence model on AQI history")
bullets([
    "An LSTM or Transformer trained on historical AQI has no causal handle on "
    "emissions. It cannot answer \u201cwhat if construction halts at T+6\u201d \u2014 "
    "the counterfactual is structurally outside its hypothesis space.",
    "It does not represent the coupling the PS is explicitly about, so it fails "
    "the brief even if its error metrics are respectable.",
    "It is what most competing teams will build.",
])

h3("Selected: physics-constrained coupled emulator")

para("A neural emulator of a coupled reference model, with the feedback terms "
     "carried as explicit, inspectable state rather than hidden weights.")

table(
    ["Design choice", "What it is", "Why this and not the alternative"],
    [
        ["Backbone",
         "Convolutional encoder\u2013decoder (U-Net shape) over the NCR grid, "
         "with a recurrent step advancing one hour at a time for 72 steps.",
         "The state is a spatial field on a regular grid and the physics is "
         "local advection and diffusion \u2014 exactly what convolution encodes. "
         "A graph network buys nothing on a regular grid; a plain MLP throws away "
         "spatial structure. Autoregressive hourly stepping matches how the "
         "physical model itself integrates, so the emulator learns a transition "
         "operator rather than a 72-hour black box."],
        ["Explicit feedback channels",
         "PBL height, surface shortwave and aerosol load are separate output "
         "channels, not internal activations.",
         "This is what makes \u03bb computable from the forecast at all. A model "
         "that only emits AQI cannot be interrogated for regime state, and the "
         "entire differentiator disappears."],
        ["Monotonicity constraint",
         "Sign of the aerosol\u2192radiation and radiation\u2192PBL responses "
         "constrained during training via a penalty term.",
         "Prevents the emulator learning a physically inverted feedback that "
         "happens to fit the training set. This is what \u201cphysics-informed\u201d "
         "means concretely, rather than as a slogan."],
        ["Loss",
         "Weighted MSE in log space, plus a term penalising error in the timing "
         "of threshold crossings.",
         "Log space because PM2.5 is log-normally distributed and plain MSE would "
         "be dominated by the few extreme hours. The timing term exists because "
         "for Disaster Management, being right about WHEN matters as much as "
         "being right about HOW MUCH."],
        ["Uncertainty",
         "Small ensemble from perturbed initial conditions plus MC dropout; "
         "spread reported alongside every forecast.",
         "Cheaper than a full ensemble system, and it lets the system report when "
         "not to trust itself \u2014 which for a disaster-management product is "
         "worth more than a confident wrong number."],
        ["Training reference",
         "A coupled reference dataset \u2014 WRF-Chem runs over historical Delhi "
         "episodes, and/or open coupled products (CAMS, SILAM, GEOS-CF) as both "
         "boundary conditions and the baseline to beat.",
         "Deliberately not fixed yet. Committing to a training source before "
         "confirming the data can actually be obtained is how projects die in "
         "week five. Module B settles it in week two."],
    ],
    [3.1, 5.4, 7.5],
)

h2("5.2  Why every module is written as functions")

para("This is not a style preference. Three specific properties are being bought:")

bullets([
    [("Testability of the scientific claim.", True), " log_anomaly(), ols_slope(), "
     "auc_mann_whitney() and find_runs() each have one job and a checkable "
     "output. An off-by-one in find_runs() silently shifts every episode "
     "boundary by an hour and corrupts the lead-time result \u2014 as a function "
     "it can be tested against hand-built cases; inline in a loop it cannot."],
    [("Substitutability at the boundaries.", True), " The training reference is "
     "undecided and the PBL source may change from ERA5 to ceilometer-corrected. "
     "Because loading, deriving and estimating are separate functions, swapping "
     "the source touches one function, not the pipeline."],
    [("Auditability by a reviewer.", True), " A judge who asks \u201chow did you "
     "estimate that derivative\u201d can be shown a fifteen-line function with no "
     "hidden regularisation. That is why ordinary least squares is hand-written "
     "rather than imported \u2014 there is visibly nothing else changing the "
     "answer."],
])

h2("5.3  Why Pathway")

para("Pathway is already load-bearing in the existing codebase, and it earns its "
     "place here for a reason that is specific to this problem rather than "
     "generic streaming enthusiasm.")

table(
    ["Property", "Why it matters for PS 26082"],
    [
        ["Event-time windows with late and out-of-order data",
         "CPCB stations drop out and backfill. Satellite AOD arrives 30 minutes "
         "to 3 hours after overpass. FIRMS NRT lands within 3 hours. A "
         "processing-time system would either wait for the slowest source or "
         "silently compute on incomplete windows. Event-time semantics let each "
         "source arrive when it arrives and still land in the correct hour."],
        ["Incremental recomputation",
         "When a late observation arrives, only the affected downstream state is "
         "recomputed rather than the whole window. This is what makes continuous "
         "re-anchoring of a 72-hour forecast affordable."],
        ["Replay determinism",
         "The same input sequence produces the same output sequence. This is the "
         "foundation of the audit trail: an escalation decision can be replayed "
         "line by line months later. For a regulatory system, this is not a "
         "convenience, it is the product."],
        ["Continuous re-anchoring \u2014 the strategic point",
         "NCMRWF runs operational cycles on a fixed HPC schedule. A streaming "
         "system can re-anchor its forecast against new observations far more "
         "often. This is the one axis on which a student team can legitimately "
         "outperform the institution that wrote the problem statement \u2014 not "
         "model physics, but update cadence and delivery."],
    ],
    [4.4, 11.2],
)

callout("the honest caveat",
        "Pathway is the right tool for ingestion, windowing, assimilation "
        "triggering and the decision layer. It is not a numerical model runner. "
        "The emulator is invoked from the DAG; it does not live inside it.")

# ================================================================= 6

h1("Target architecture", "6")

code("""
================================ AREE 2.0 ================================

  DATASET A (real-time)                    DATASET B (2021-2025 archive)
  CPCB/OpenAQ  GFS/Open-Meteo              CPCB archive  ERA5  MODIS/MAIAC
  INSAT-3D AOD  FIRMS NRT                  FIRMS SP  CAMS-GLOB-ANT
        |                                        |
        v                                        v
  +-------------------+                  +-------------------+
  | MODULE A          |                  | MODULE A          |
  | Ingestion         |                  | Historical build  |
  | Pathway connectors|                  | one hourly panel  |
  | observed_at /     |                  +---------+---------+
  | received_at /     |                            |
  | processed_at      |                            v
  +---------+---------+                  +-------------------+
            |                            | MODULE B          |
            v                            | lambda estimation |
  +-------------------+                  | elasticities, IV  |
  | OBSERVATIONAL     |                  | episode labelling |
  | ANCHORING         |                  | validation gate   |
  | nudging / OI      |                  +---------+---------+
  | (NOT 3D-Var)      |                            |
  +---------+---------+                            v
            |                            +-------------------+
            |                            | MODULE C          |
            +--------------------------->| Coupled emulator  |
                                         | 72 h, hourly, grid|
                                         | explicit feedback |
                                         | channels          |
                                         +---------+---------+
                                                   |
                                                   v
                                         +-------------------+
                                         | MODULE B (online) |
                                         | REGIME DIAGNOSTIC |
                                         | lambda(t)         |
                                         | lock-in time      |
                                         | uncertainty       |
                                         +---------+---------+
                                                   |
                                                   v
                                         +-------------------+
                                         | INTERVENTION      |
                                         | WINDOW +          |
                                         | COST THRESHOLD    |
                                         +---------+---------+
                                                   |
                                                   v
                                         +-------------------+
                                         | MODULE D - AREE   |
                                         | deterministic rule|
                                         | engine (reused)   |
                                         | authority mapping |
                                         | escalation, audit |
                                         +---------+---------+
                                                   |
                                                   v
                                         +-------------------+
                                         | MODULE E          |
                                         | Next.js ops UI    |
                                         | forecast, regime, |
                                         | window, orders    |
                                         +-------------------+
""")

h2("The layer that was missing from the previous sketch")

callout("false-alarm cost",
        "GRAP Stage III has real economic cost \u2014 construction halts, school "
        "closures. So the question \u201cat what forecast probability do you "
        "actually act?\u201d is a decision-threshold problem with asymmetric "
        "costs, and it was absent from the earlier architecture. It belongs "
        "between the intervention window and the AREE core, and it is the best "
        "possible home for the deterministic engine: the threshold is encoded as "
        "policy, not decided by a model, and the audit trail records which "
        "threshold was in force when an order fired. The \u201cAI explains, AI "
        "does not decide\u201d principle extends into it cleanly.", color=BLUE)

# ================================================================= 7

h1("Module specification and team allocation", "7")

para("Five owners. Module boundaries are drawn so that each can be developed and "
     "tested without waiting on another, and so that the interface between any "
     "two is a file on disk or a typed dictionary \u2014 never shared mutable "
     "state.")

for mod, owner, mission, deliv, iface, done in [
    ("MODULE A", "Data engineering",
     "Own every byte that enters the system, and the honesty of its timestamps.",
     ["ERA5 pull for the NCR domain, 2021\u20132025, Oct\u2013Feb "
      "(scripts/01_fetch_era5.py \u2014 written, needs a CDS key).",
      "Ground PM2.5 via OpenAQ v3 or the S3 archive (scripts/02).",
      "MODIS MCD19A2 AOD via Google Earth Engine; INSAT-3D AOD via MOSDAC.",
      "FIRMS NRT + SP ingestion \u2014 extend the existing firms_stream.py.",
      "Fill in the empty weather_stream.py with a real met connector.",
      "Three-timestamp discipline everywhere: observed_at, received_at, "
      "processed_at, and a data-age field surfaced to the UI."],
     "data/processed/panel_hourly.parquet, and live Pathway tables with the "
     "same column names.",
     "The panel builds end-to-end, ERA5 accumulations are correctly converted "
     "to W m\u207b\u00b2, and no step in the pipeline can tell whether it is "
     "reading live or replayed data."),

    ("MODULE B", "Feedback science",
     "Own \u03bb: its estimation, its validation, and the honesty of the claim.",
     ["Run the validation gate on real data and report the held-out AUC "
      "(scripts/04\u201306 \u2014 written and self-tested).",
      "Fix the simultaneity bias with instrumental-variable estimation; "
      "synoptic wind is the natural instrument.",
      "Quantify ERA5 PBLH bias against IMD radiosonde over Delhi, and against "
      "ceilometer if a series can be obtained.",
      "Define the regime classification bands and the lock-in time estimator.",
      "Uncertainty conditioned on \u03bb: does forecast spread widen as \u03bb "
      "approaches 1?"],
     "validation_report.json, the regime thresholds, and a written statement of "
     "what \u03bb does and does not support.",
     "A held-out AUC number exists and is defensible, or the claim is formally "
     "withdrawn and the fallback positioning adopted."),

    ("MODULE C", "Forecast emulator",
     "Own the 72-hour coupled prediction.",
     ["Settle the training reference: WRF-Chem runs, open coupled products, or "
      "a combination. Decide in week two, not week five.",
      "Build the U-Net + recurrent-step emulator with explicit PBL, shortwave "
      "and aerosol output channels.",
      "Monotonicity penalty on the feedback terms.",
      "Ensemble spread from perturbed initial conditions.",
      "Verification: MAE, RMSE, bias at 24/48/72 h against persistence and an "
      "external baseline, on the held-out Novembers.",
      "Natural-experiment validation: the 2020 lockdown emissions perturbation "
      "and the annual Diwali spike, rather than synthetic counterfactuals."],
     "A callable forecast(state) -> 72h gridded fields with uncertainty.",
     "It beats persistence at 24 h, and where it does not, that is stated "
     "plainly with the horizon at which skill is lost."),

    ("MODULE D", "Regulatory engine",
     "Own the conversion of physical risk into accountable action.",
     ["Reuse state_machine.py essentially as-is \u2014 it is already correct.",
      "Retune PERSISTENCE_THRESHOLD from 3 windows of 3 minutes to a "
      "meteorologically meaningful horizon.",
      "Add the predicted-breach path alongside the observed-breach path.",
      "Build the false-alarm cost threshold layer.",
      "Jurisdiction \u2192 authority mapping; escalation case with deadline.",
      "Hash-chained audit; signed order export via report_generator.py."],
     "Escalation cases with rule citation, authority, deadline and a replayable "
     "decision record.",
     "Any escalation can be replayed from the audit log and reproduces "
     "byte-identically."),

    ("MODULE E", "Interface and integration",
     "Own what a human actually sees, and the seams between modules.",
     ["Forecast view: 72 h, gridded, on the existing Leaflet map.",
      "Regime view: \u03bb through time with the threshold and the projected "
      "lock-in hour.",
      "Intervention window as the primary call to action, not a number.",
      "Data-age and provenance surfaced on every panel.",
      "Keep Next.js free of business logic \u2014 the existing discipline is "
      "correct and must hold.",
      "Docker compose, environment handling, end-to-end demo script."],
     "A running system and a rehearsed demo path.",
     "The demo runs twice in a row without manual intervention."),
]:
    h2(f"{mod}  \u2014  {owner}")
    para([("Mission. ", True), (mission, False)])
    h3("Deliverables")
    bullets(deliv)
    h3("Interface out")
    para(iface, indent=0.55, after=4)
    h3("Definition of done")
    para(done, indent=0.55, after=4)

h2("The fifth seat: presentation")
para("The PPT owner is not a documentation clerk. They own the argument: the "
     "one-line positioning, the lock-in demo narrative, the verification slide, "
     "and the answers to the three questions NCMRWF will certainly ask \u2014 "
     "what is your skill score, how did you handle PBLH uncertainty, and what "
     "does a false alarm cost. That work starts now, not in the last week, "
     "because it determines which results the technical modules must produce.")

# ================================================================= 8

h1("Evaluation protocol", "8")

para("A forecasting problem statement judged by a forecasting centre will be "
     "asked for skill scores. A team with honest verification beats a team with "
     "a prettier dashboard in front of these judges, every time.")

table(
    ["Category", "Metric", "Compared against", "Reported at"],
    [
        ["Forecast skill", "MAE, RMSE, mean bias on PM2.5",
         "Persistence, and an external coupled product (CAMS or similar)",
         "24 h, 48 h, 72 h"],
        ["Regime detection", "Held-out AUC, hit rate, false-alarm rate",
         "Negative ventilation coefficient", "Per episode"],
        ["Lead time", "Median hours from alert to onset", "\u2014",
         "Per locked-in episode"],
        ["Uncertainty", "Interval coverage, spread\u2013error correlation",
         "Nominal coverage", "All horizons"],
        ["Counterfactual", "Reproduction of an observed emissions perturbation",
         "2020 lockdown, Diwali spikes", "Case study"],
        ["Decision layer", "Response latency", "\u2014",
         "Separately, and never called forecast latency"],
    ],
    [2.8, 4.6, 5.2, 3.0],
)

callout("hold-out discipline",
        "November 2023 and November 2024 are excluded from every fitting step in "
        "the pipeline as written. Those two months are the most severe recent "
        "post-monsoon episodes. Any number quoted publicly comes from them.",
        color=BLUE)

# ================================================================= 9

h1("Risk register", "9")

table(
    ["Risk", "Severity", "Mitigation, and who owns it"],
    [
        ["ERA5 PBLH is biased in exactly the shallow stable layers that matter "
         "most. Known to struggle with nocturnal and heavily-loaded cases.",
         "High",
         "Module B quantifies the bias against radiosonde rather than ignoring "
         "it, and the emulator is allowed to learn a correction. Stated openly "
         "in the presentation."],
        ["\u03bb does not separate the classes on real data.",
         "High",
         "This is why the gate exists in week one. Fallback positioning is "
         "already written: coupled forecasting with an explicit feedback "
         "diagnostic, no lock-in claim."],
        ["Simultaneity bias in the elasticity estimates.",
         "Medium",
         "Documented in section 3.5. Module B replaces OLS with instrumental "
         "variables."],
        ["Training reference for the emulator cannot be obtained in time.",
         "High",
         "Do not architect around a specific model before confirming data "
         "access. Module C decides in week two with a written go/no-go."],
        ["CPCB dry mass versus satellite wet AOD are not the same quantity.",
         "Medium",
         "Hygroscopic growth correction is required the moment AOD enters. "
         "Module A flags it; Module C implements it."],
        ["Scope inflation \u2014 the architecture is larger than five people.",
         "High",
         "Module boundaries are drawn to be independently shippable. If a module "
         "slips, the others still demo. Cut Module C's spatial resolution before "
         "cutting verification."],
        ["Presenting the 23 ms latency as forecast speed.",
         "Medium",
         "Removed from all material. Module E owns compliance."],
    ],
    [5.0, 1.9, 8.7],
    colors={1: RED},
)

# ================================================================= 10

h1("What to run first", "10")

para("Steps one to three are one person for about a week and they either prove "
     "or kill the central idea. Nothing else should start on the assumption that "
     "\u03bb works.")

code("""
    cd d:\\AZUO\\chat\\aree-ps26082

    # already passing - no keys needed, proves the estimator works
    ..\\.venv\\Scripts\\python.exe scripts\\00_selftest_synthetic.py

    # 1. confirm CDS access, then pull ERA5   <-- GATES EVERYTHING
    set CDSAPI_KEY=<key>
    ..\\.venv\\Scripts\\python.exe scripts\\01_fetch_era5.py --dry-run
    ..\\.venv\\Scripts\\python.exe scripts\\01_fetch_era5.py

    # 2. ground PM2.5
    set OPENAQ_API_KEY=<key>
    ..\\.venv\\Scripts\\python.exe scripts\\02_fetch_ground_aq.py

    # 3. build the panel, estimate lambda, label episodes, validate
    ..\\.venv\\Scripts\\python.exe scripts\\03_build_panel.py
    ..\\.venv\\Scripts\\python.exe scripts\\04_compute_lambda.py
    ..\\.venv\\Scripts\\python.exe scripts\\05_label_episodes.py
    ..\\.venv\\Scripts\\python.exe scripts\\06_validate_lambda.py
""")

table(
    ["Order", "Task", "Owner", "Gate"],
    [
        ["1", "Confirm Copernicus CDS access and pull ERA5 blh, ssrd, ssrdc for "
         "five winters.", "Module A",
         "If PBLH cannot be obtained at hourly cadence, \u03bb as formulated is "
         "dead. Know this first."],
        ["2", "Ground PM2.5 for the same hours.", "Module A", "\u2014"],
        ["3", "Run the validation gate. Report held-out AUC and lead time.",
         "Module B", "Decides whether the lock-in claim is made at all."],
        ["4", "Settle the emulator training reference.", "Module C",
         "Written go/no-go on data availability before any model code."],
        ["5", "Retune the persistence constants and add the predicted-breach "
         "path.", "Module D", "\u2014"],
        ["6", "Only now: build the emulator.", "Module C", "\u2014"],
    ],
    [1.4, 7.2, 2.4, 5.6],
)

P(after=14)
rule(color="C8CFD8", size=4, before=10, after=10)
p = P(after=0, line=1.3)
R(p, "The honest summary: the architecture is sound, the estimator is written "
     "and passes its self-test, and the entire scientific claim now rests on one "
     "week of data work that has not yet been done. That is the right place for "
     "the risk to sit.", size=10.5, italic=True, color=INK, font=SERIF)

p = P(before=10, after=0)
R(p, "Team Devengers   \u00b7   PS 26082   \u00b7   Architecture revision A",
  size=7.5, font=MONO, color=MUTE, caps=True, space=0.6)

doc.save(OUT)
print("saved:", OUT)
