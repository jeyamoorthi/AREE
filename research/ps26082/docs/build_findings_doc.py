# -*- coding: utf-8 -*-
"""Build the PS 26082 validation findings document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
OUT = HERE / "AREE_PS26082_Findings.docx"

INK = RGBColor(0x14, 0x18, 0x20)
BODY = RGBColor(0x33, 0x3C, 0x4A)
MUTE = RGBColor(0x77, 0x81, 0x90)
BLUE = RGBColor(0x14, 0x47, 0x6E)
RED = RGBColor(0x9E, 0x2B, 0x25)
GREEN = RGBColor(0x1E, 0x6B, 0x44)

SERIF, SANS, MONO = "Georgia", "Segoe UI", "Consolas"

doc = Document()


def _rf(run, font):
    rPr = run._element.get_or_add_rPr()
    rf = rPr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), font)


def R(p, text, size=10, bold=False, italic=False, color=BODY, font=SANS,
      caps=False, space=None):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.all_caps = caps
    _rf(r, font)
    if space is not None:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(space * 20)))
        r._element.get_or_add_rPr().append(sp)
    return r


def P(before=0, after=6, line=1.25, indent=None, keep=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after, pf.line_spacing = Pt(before), Pt(after), line
    pf.keep_with_next = keep
    if indent is not None:
        pf.left_indent = Cm(indent)
    return p


def para(parts, **kw):
    p = P(**kw)
    if isinstance(parts, str):
        R(p, parts)
    else:
        for t, b in parts:
            R(p, t, bold=b, color=INK if b else BODY)
    return p


def rule(color="C8CFD8", size=6, before=0, after=8):
    p = P(before=before, after=after, line=1.0)
    bdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), color)
    bdr.append(b)
    p._p.get_or_add_pPr().append(bdr)


def h1(text, num=None, newpage=True):
    if newpage:
        doc.add_page_break()
    p = P(before=0, after=3, keep=True)
    if num:
        R(p, f"{num}   ", size=19, bold=True, color=BLUE, font=SERIF)
    R(p, text, size=19, bold=True, color=INK, font=SERIF)
    rule("14476E", 14, after=10)


def h2(text):
    p = P(before=13, after=4, keep=True)
    R(p, text, size=12, bold=True, color=INK, font=SERIF)


def h3(text):
    p = P(before=9, after=3, keep=True)
    R(p, text, size=9, bold=True, color=BLUE, font=MONO, caps=True, space=0.8)


def bullets(items):
    for it in items:
        p = P(after=3, line=1.2, indent=0.55)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        R(p, "•   ", size=9, color=BLUE, font=MONO)
        if isinstance(it, str):
            R(p, it, size=9.5)
        else:
            for t, b in it:
                R(p, t, size=9.5, bold=b, color=INK if b else BODY)


def code(text, size=8.2):
    for line in text.strip("\n").split("\n"):
        p = P(after=0, line=1.02, indent=0.35)
        R(p, line if line else " ", size=size, font=MONO, color=INK)
    P(after=6, line=1.0)


def callout(title, text, color=RED):
    p = P(before=8, after=2, line=1.2, indent=0.3, keep=True)
    R(p, title, size=9, bold=True, font=MONO, caps=True, color=color, space=0.7)
    p = P(after=8, line=1.22, indent=0.3)
    if isinstance(text, str):
        R(p, text, size=9.5)
    else:
        for t, b in text:
            R(p, t, size=9.5, bold=b, color=INK if b else BODY)


def cell_border(cell, **kw):
    tcPr = cell._tc.get_or_add_tcPr()
    bs = tcPr.find(qn("w:tcBorders"))
    if bs is None:
        bs = OxmlElement("w:tcBorders")
        tcPr.append(bs)
    for edge in ("top", "bottom"):
        spec = kw.get(edge)
        if spec is None:
            continue
        el = bs.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            bs.append(el)
        sz, col = spec
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), col)


def table(headers, rows, widths, fs=8.5, colcolors=None):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tblPr = t._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    nb = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{e}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        nb.append(el)
    tblPr.append(nb)
    mar = OxmlElement("w:tblCellMar")
    for e, v in (("top", 55), ("left", 0), ("bottom", 55), ("right", 130)):
        el = OxmlElement(f"w:{e}")
        el.set(qn("w:w"), str(v))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)
    t.autofit = False
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)
    for row in t.rows:
        for c, w in zip(row.cells, widths):
            c.width = Cm(w)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_border(c, bottom=(10, "9AA5B2"))
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        R(p, h, size=7.5, font=MONO, color=MUTE, caps=True, space=0.7)
    for ri, row in enumerate(rows, 1):
        for ci, v in enumerate(row):
            c = t.rows[ri].cells[ci]
            if ri > 1:
                cell_border(c, top=(4, "DDE2E9"))
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.12
            col = (colcolors or {}).get(ci, BODY)
            R(p, str(v), size=fs, color=col, bold=(ci == 0),
              font=MONO if ci and str(v).replace(".", "").replace("-", "").isdigit() else SANS)
    P(after=4)


# ---------------------------------------------------------------- setup
st = doc.styles["Normal"]
st.font.name = SANS
st.font.size = Pt(10)
st.font.color.rgb = BODY
st.element.rPr.rFonts.set(qn("w:eastAsia"), SANS)
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.line_spacing = 1.25

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(1.8)
sec.left_margin, sec.right_margin = Cm(2.2), Cm(2.2)

f = sec.footer.paragraphs[0]
f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
R(f, "PS 26082 validation findings  ·  Team Devengers  ·  page ",
  size=7, font=MONO, color=MUTE)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
f._p.append(fld)

# ---------------------------------------------------------------- cover
p = P(after=4, line=1.0)
R(p, "PROBLEM STATEMENT 26082  ·  VALIDATION FINDINGS", size=8, font=MONO,
  color=BLUE, bold=True, caps=True, space=1.5)
p = P(after=12, line=1.0)
R(p, "Ministry of Earth Sciences  ·  NCMRWF  ·  Team Devengers",
  size=8, font=MONO, color=MUTE, caps=True, space=1.3)

p = P(after=3, line=1.0)
R(p, "The λ hypothesis is falsified.", size=26, bold=True, color=INK, font=SERIF)
p = P(after=8, line=1.2)
R(p, "What five winters of Delhi NCR data actually say, and where the "
     "project goes now.", size=12.5, italic=True, color=BODY, font=SERIF)
rule("9E2B25", 18, before=4, after=10)

para([("The aerosol–radiation–boundary-layer feedback loop gain "
       "λ does not separate pollution episodes that lock in from those "
       "that ventilate out. Held-out AUC is ", False),
      ("0.504", True), (" — chance. The claim is withdrawn, and no part "
       "of the submission should rest on it.", False)])

P(after=6)
callout("this is the gate working, not the project failing",
        [("The validation gate existed precisely to find this out in week "
          "one rather than week six. It cost one day of data engineering. "
          "The alternative was five people building a presentation around a "
          "number that would have collapsed under the first question from an "
          "NCMRWF reviewer.", False)], color=GREEN)

h3("What replaced it")
para([("Lock-in is not diagnosable from the present state. It is determined "
       "by weather that has not happened yet at episode onset. Ventilation "
       "measured ", False), ("before", True), (" onset has no skill (AUC "
       "0.514); measured ", False), ("after", True), (" onset it is skilful "
       "(AUC 0.736). That is a forecasting problem, which is exactly what PS "
       "26082 asks for — so the project stays on-brief without the "
       "λ claim.", False)])

# ---------------------------------------------------------------- 1
h1("What was built and run", "1")

para("Everything below is from real data, executed end to end. Nothing is "
     "projected or assumed.")

table(
    ["Component", "What was obtained"],
    [
        ["Meteorology",
         "ERA5 via the Open-Meteo archive, 9-point grid across the NCR box, "
         "hourly 2020-10-01 to 2025-03-31. 39,432 hours. Boundary layer "
         "height, surface and top-of-atmosphere shortwave, cloud, "
         "temperature, RH, wind, pressure, precipitation. No Copernicus key "
         "was needed, which removed the only hard external dependency."],
        ["Ground PM2.5",
         "34 OpenAQ sensors inside the NCR box, reference-grade networks only "
         "(CPCB, DPCC, UPPCB, HSPCB, IMD, IITM), anchored by the "
         "continuously-operating US Embassy monitor. 29,953 hours matched to "
         "the meteorology."],
        ["Episodes",
         "161 labelled from the PM2.5 trajectory alone: 39 locked in, 122 "
         "ventilated, 11 in the held-out Novembers. Labels never saw λ."],
        ["Pipeline",
         "Nine scripts, from synthetic self-test through elasticity "
         "estimation, episode labelling, validation, diagnosis and the "
         "reformulation attempt. All committed on branch "
         "ps26082-coupled-forecast."],
    ],
    [3.4, 13.2],
)

h2("A data limitation that must be stated")

callout("OpenAQ has a hole in the Indian record",
        "The old CPCB sensor generation stops around November 2022; the new "
        "generation begins February 2025. Nothing in between except the US "
        "Embassy monitor. Consequently 64% of in-season hours rest on two or "
        "fewer stations, including both holdout Novembers. The domain median "
        "is effectively one instrument across that stretch. This weakens "
        "spatial representativeness and is one of the candidate explanations "
        "examined in section 3.")

table(
    ["Season", "Hours", "Median stations", "PM2.5 p95"],
    [
        ["2020-21", "3,171", "15", "365"],
        ["2021-22", "3,209", "21", "322"],
        ["2022-23", "2,071", "1", "334"],
        ["2023-24", "2,603", "1", "359"],
        ["2024-25", "3,375", "1", "347"],
    ],
    [3.0, 3.0, 4.6, 4.0],
)

# ---------------------------------------------------------------- 2
h1("The result", "2")

h2("2.1  The mechanism is unmistakably present in the raw data")

para("Before anything failed, this succeeded. The median diurnal cycle over "
     "five in-season winters is exactly what the physics predicts:")

code("""
    hour IST     PM2.5 (ug/m3)     BLH (m)     clearness
    ---------------------------------------------------
       23             163              44          -
       06             152              52          -
       09             151             196        0.554
       12             104             863        0.660
       15              79            1054        0.586
       18             116              60          -
""")

para([("PM2.5 falls by half as the mixed layer deepens twenty-five fold, and "
       "recovers as it collapses. The aerosol–boundary-layer relationship "
       "is not in doubt. ", False),
      ("What failed is the attempt to extract a loop GAIN from it.", True)])

h2("2.2  λ returns chance")

table(
    ["Predictor", "AUC (train)", "Hit / false alarm", "Verdict"],
    [
        ["λ", "0.504", "0.86 / 0.73", "chance"],
        ["negative ventilation coefficient", "0.421", "0.42 / 0.33", "worse than chance"],
        ["ventilation-failure index (step 8)", "0.554", "0.31 / 0.17", "chance"],
        ["seasonal-anomaly ventilation index", "0.512", "0.62 / 0.52", "chance"],
    ],
    [6.0, 3.0, 3.8, 3.8],
    colcolors={3: RED},
)

para("A hit rate of 0.86 alongside a false-alarm rate of 0.73 is not a "
     "detector. It is a threshold low enough to fire on almost everything. "
     "The AUC is the honest summary and it is 0.504.")

# ---------------------------------------------------------------- 3
h1("Why it failed", "3")

para("Three candidate explanations were tested rather than assumed. The "
     "diagnostic script is committed as 07_diagnose_elasticities.py.")

h2("3.1  Errors-in-variables — the dominant cause")

para("λ requires differentiating ERA5 boundary-layer-height anomalies. "
     "Classical measurement error in a regressor biases the slope toward "
     "zero by the reliability ratio. Reverse regression brackets it:")

code("""
    forward   d lnC / d lnH          = -0.208    (attenuated)
    reverse   1 / (d lnH / d lnC)    = -12.743   (over-stated)
    implied reliability ratio          0.016
""")

para([("A reliability ratio of ", False), ("0.016", True), (" means the ERA5 "
       "BLH anomaly signal is roughly 98% noise once the diurnal cycle is "
       "removed. ERA5 derives boundary layer height from a bulk Richardson "
       "number and is known to struggle in shallow stable layers — which "
       "is the entire Delhi winter night. The elasticity is not small because "
       "the physics is weak; it is small because the instrument is.", False)])

h2("3.2  The sign is not stable")

table(
    ["Specification", "e1 = dlnC/dlnH", "Expected", "Status"],
    [
        ["Pooled hourly", "-0.043", "≈ -1", "right sign, tiny"],
        ["Dense network (>= 8 stations)", "+0.061", "≈ -1", "WRONG SIGN"],
        ["Daily daytime aggregates", "+0.084", "≈ -1", "WRONG SIGN"],
    ],
    [5.4, 3.6, 2.6, 4.6],
    colcolors={3: RED},
)

callout("the synthetic test passed and the real data did not — why",
        "The synthetic generator built the loop in by construction, so the "
        "estimator had a clean signal with modest noise and recovered "
        "λ to 9%. That validated the ARITHMETIC, not the premise that "
        "the signal survives in ERA5. A self-test can only falsify an "
        "estimator, never confirm that real data contains what you hope. "
        "This is worth saying out loud in the presentation: it is the "
        "difference between a team that tested its method and one that "
        "believed it.")

h2("3.3  A reformulation avoiding PBL height entirely — also fails")

para("If ERA5 BLH is the problem, remove it. Step 8 defines a ventilation "
     "index from observed PM2.5 alone:")

code("""
    V = ln( PM2.5 pre-dawn  /  PM2.5 afternoon )
""")

para("Large V means the mixed layer grew and cleared the overnight load. V "
     "near zero means it did not. It uses only a reference-grade BAM "
     "measurement and no derivative of any model field.")

bullets([
    "V behaves correctly as a ventilation measure: correlation +0.21 with "
    "daytime maximum BLH, +0.15 with mean wind.",
    "15% of in-season days show V <= 0 — no net daytime dilution at all.",
    [("But it does not predict lock-in: AUC 0.554.", True)],
])

para("So the failure is not an artefact of ERA5 alone. Something more "
     "fundamental is going on, and section 4 identifies it.")

# ---------------------------------------------------------------- 4
h1("The finding that redirects the project", "4")

para("If lock-in cannot be predicted from the state before onset, the "
     "obvious question is whether it is predictable from anything. It is — "
     "from the weather that follows.")

table(
    ["Predictor", "Window", "AUC", "Verdict"],
    [
        ["Mean wind speed", "-24 to 0 h (before onset)", "0.434", "no skill"],
        ["Mean wind speed", "0 to +48 h (after onset)", "0.672", "weak"],
        ["Daytime max BLH", "-24 to 0 h (before onset)", "0.567", "no skill"],
        ["Daytime max BLH", "0 to +48 h (after onset)", "0.724", "SKILFUL"],
        ["Ventilation coefficient", "-24 to 0 h (before onset)", "0.514", "no skill"],
        ["Ventilation coefficient", "0 to +48 h (after onset)", "0.736", "SKILFUL"],
    ],
    [4.6, 5.0, 2.4, 4.2],
    colcolors={3: INK},
)

h2("What this means")

para([("Whether a Delhi pollution episode lasts sixteen hours or ten days is "
       "decided by the ventilation over the following two days — not by "
       "the aerosol–boundary-layer state at the moment it begins. ", False),
      ("Lock-in is not a diagnosable regime. It is a forecastable outcome.",
       True)])

para("That is a genuine scientific finding, it is measured rather than "
     "asserted, and it lands the project exactly where the problem statement "
     "wanted it: on coupled forecasting.")

h2("The revised positioning")

p = P(before=6, after=8, indent=0.3, line=1.3)
R(p, "We measured whether pollution lock-in can be diagnosed from the "
     "present state. It cannot. It is set by the ventilation over the next "
     "48 hours — so we forecast that, and convert it into the time "
     "remaining to act.", size=12, italic=True, color=BLUE, font=SERIF)

para("This is stronger than the original pitch in three ways. It is backed "
     "by a measurement rather than a hypothesis. It explains why a coupled "
     "forecast is necessary rather than merely desirable — ventilation is "
     "the quantity that decides the outcome, and it depends on the "
     "meteorology the PS says must be coupled. And it survives the question "
     "that would have destroyed the λ pitch: a reviewer asking what the "
     "held-out skill score is.")

# ---------------------------------------------------------------- 5
h1("What changes for each module", "5")

table(
    ["Module", "Change"],
    [
        ["A — Data",
         "Unchanged and validated. The Open-Meteo path removed the CDS "
         "dependency entirely. Two new connectors are committed: "
         "cpcb_stream.py (authoritative CPCB via data.gov.in, with measured "
         "data age) and weather_stream.py, which was 0 bytes and is now a "
         "working 72-hour meteorological feed. Priority: close the OpenAQ "
         "2022-2025 gap from CPCB directly so the holdout rests on more than "
         "one instrument."],
        ["B — Feedback science",
         "The λ mandate is closed. Reassign to the ventilation forecast: "
         "define the operational ventilation threshold, quantify how far "
         "ahead it can be skilfully forecast, and own the false-alarm cost "
         "curve. The diagnostic work in scripts 07 and 08 is the evidence "
         "base and should be presented, not hidden."],
        ["C — Emulator",
         "Now the centre of the project, with a sharper target. The emulator "
         "must forecast the ventilation coefficient and boundary-layer "
         "evolution well, because those are the variables shown to determine "
         "the outcome. Verification against persistence at 24/48/72 h is "
         "still the number that decides the score."],
        ["D — Regulatory engine",
         "Unchanged in design, better justified. It now converts a forecast "
         "ventilation collapse into lead time and an escalation case. The "
         "false-alarm cost threshold matters more than before, because the "
         "decision now rests on a forecast rather than an observation."],
        ["E — Interface",
         "The primary display becomes the 72-hour ventilation forecast and "
         "the intervention window, not a λ gauge. Data age and station "
         "count must be visible — 64% of the historical record rests on "
         "two or fewer stations and the interface should never imply "
         "otherwise."],
    ],
    [3.0, 13.6],
)

# ---------------------------------------------------------------- 6
h1("What to do next", "6")

table(
    ["Order", "Task", "Owner"],
    [
        ["1", "Close the OpenAQ gap: pull CPCB 2022-2025 directly so the "
              "holdout rests on a network rather than one monitor.", "A"],
        ["2", "Establish how far ahead the ventilation coefficient is "
              "skilfully forecastable, using Open-Meteo forecasts verified "
              "against the ERA5 archive. This sets the honest lead-time "
              "claim.", "B"],
        ["3", "Settle the emulator training reference with a written "
              "go/no-go on data availability.", "C"],
        ["4", "Retune the persistence constants from 3 windows of 3 minutes "
              "to a meteorologically meaningful horizon; add the "
              "predicted-breach path.", "D"],
        ["5", "Rebuild the pitch around the measured finding. The λ "
              "slides are withdrawn.", "PPT"],
    ],
    [1.5, 11.5, 1.8],
)

P(after=10)
rule("C8CFD8", 4, before=8, after=10)
p = P(after=0, line=1.3)
R(p, "A negative result obtained in week one, with the evidence to explain "
     "it and a measured finding to replace it, is a better position than an "
     "unexamined hypothesis carried to the finale.",
  size=10.5, italic=True, color=INK, font=SERIF)

p = P(before=8, after=0)
R(p, "Team Devengers   ·   PS 26082   ·   Findings revision A",
  size=7.5, font=MONO, color=MUTE, caps=True, space=0.6)

doc.save(OUT)
print("saved:", OUT)
