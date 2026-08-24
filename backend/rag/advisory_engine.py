# Policy-grounded advisory engine
# Pathway DocumentStore with live streaming re-indexing
# policies/ -> parse -> chunk -> embed -> BruteForceKnn retrieval

import io
import json as _json
import os
import threading
import numpy as np
from datetime import datetime, timezone

import pathway as pw
from pathway.xpacks.llm.document_store import DocumentStore
from pathway.xpacks.llm.embedders import SentenceTransformerEmbedder
from pathway.xpacks.llm.splitters import TokenCountSplitter
from pathway.stdlib.indexing import BruteForceKnnFactory
from sentence_transformers import SentenceTransformer

from config import POLICY_DIR, PERSISTENCE_THRESHOLD, HIGH_AQI_THRESHOLD

os.makedirs(POLICY_DIR, exist_ok=True)

# shared state for the UI
_rag_state = {
    "index_type": "Pathway DocumentStore (Live Hybrid Index)",
    "docs_indexed": 0,
    "chunks_indexed": 0,
    "embed_model": "all-MiniLM-L6-v2",
    "last_reindex": None,
    "store_status": "starting",
    "policy_files": [],
    "error": None,
    # Per-file extraction outcome, so an unparseable document is visible in the
    # console instead of silently missing from the index.
    "parse_errors": [],
}


# --- Document text extraction -------------------------------------------
# The policy folder holds .txt, .pdf and .docx. Reading everything as plaintext
# made Pathway's connector fail UTF-8 decoding on the PDF, and the PDF's text
# never reached the index at all. Each format now gets its proper parser.

SUPPORTED_EXTENSIONS = {".txt", ".md", ".text", ".pdf", ".docx"}


def _extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx_text(data: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _record_parse_error(name, message):
    entry = {"file": name, "error": message}
    errors = [e for e in _rag_state["parse_errors"] if e["file"] != name]
    errors.append(entry)
    _rag_state["parse_errors"] = errors
    print(f"[RAG] parse error {name}: {message}")


def _clear_parse_error(name):
    if any(e["file"] == name for e in _rag_state["parse_errors"]):
        _rag_state["parse_errors"] = [
            e for e in _rag_state["parse_errors"] if e["file"] != name
        ]


def extract_document_text(path, data):
    """Extract plain text from a policy document by file extension.

    Returns "" for unsupported or unparseable files rather than raising, so one
    bad document cannot stall the streaming index.
    """
    name = os.path.basename(path or "")
    ext = os.path.splitext(name)[1].lower()

    if not data:
        return ""

    try:
        if ext == ".pdf":
            text = _extract_pdf_text(data)
        elif ext == ".docx":
            text = _extract_docx_text(data)
        elif ext in (".txt", ".md", ".text", ""):
            text = data.decode("utf-8", errors="ignore")
        else:
            _record_parse_error(name, f"unsupported extension '{ext}'")
            return ""
    except Exception as e:  # noqa: BLE001 - never break the stream on one file
        _record_parse_error(name, f"{type(e).__name__}: {str(e)[:160]}")
        return ""

    if not text.strip():
        _record_parse_error(name, "no extractable text")
        return ""

    _clear_parse_error(name)
    return text


def _metadata_path(metadata):
    """Pull the source path out of Pathway's _metadata column."""
    if metadata is None:
        return ""
    if isinstance(metadata, dict):
        return metadata.get("path", "")
    for accessor in ("as_dict", "value"):
        candidate = getattr(metadata, accessor, None)
        if candidate is not None:
            try:
                data = candidate() if callable(candidate) else candidate
                if isinstance(data, dict):
                    return data.get("path", "")
            except Exception:  # noqa: BLE001
                pass
    try:
        return _json.loads(str(metadata)).get("path", "")
    except Exception:  # noqa: BLE001
        return ""


def _scan_policy_files():
    files = []
    for f in sorted(os.listdir(POLICY_DIR)):
        p = os.path.join(POLICY_DIR, f)
        if os.path.isfile(p):
            stat = os.stat(p)
            ext = os.path.splitext(f)[1].lower()
            error = next(
                (e["error"] for e in _rag_state["parse_errors"] if e["file"] == f), None
            )
            files.append({
                "name": f,
                "size_kb": round(stat.st_size / 1024, 1),
                "modified": datetime.fromtimestamp(
                    stat.st_mtime, tz=timezone.utc
                ).strftime("%Y-%m-%d %H:%M"),
                "type": ext.lstrip(".") or "unknown",
                "supported": ext in SUPPORTED_EXTENSIONS,
                "parse_error": error,
            })
    _rag_state["policy_files"] = files
    _rag_state["docs_indexed"] = len(files)

_scan_policy_files()


# --- Pathway pipeline setup ---

_pw_embedder = SentenceTransformerEmbedder(model="all-MiniLM-L6-v2")

# Read raw bytes rather than plaintext: the connector itself used to attempt a
# UTF-8 decode of every file and failed noisily on the PDF. Text extraction now
# happens in the DAG, per format, so downstream stages always see valid UTF-8.
_raw_policy_docs = pw.io.fs.read(
    POLICY_DIR,
    format="binary",
    mode="streaming",
    with_metadata=True,
)


@pw.udf
def _parse_policy_document(data: bytes, metadata: pw.Json) -> bytes:
    """Extract text from a policy document and re-emit it as UTF-8 bytes."""
    return extract_document_text(_metadata_path(metadata), data).encode("utf-8")


_policy_docs = _raw_policy_docs.select(
    data=_parse_policy_document(pw.this.data, pw.this._metadata),
    _metadata=pw.this._metadata,
)

_splitter = TokenCountSplitter(min_tokens=50, max_tokens=300)

_retriever_factory = BruteForceKnnFactory(
    dimensions=384,
    embedder=_pw_embedder,
)

_doc_store = DocumentStore(
    docs=_policy_docs,
    retriever_factory=_retriever_factory,
    splitter=_splitter,
)

print("[RAG] DocumentStore ready")


# --- Live index: observer captures chunks for python-side queries ---

_live_chunks = {}
_live_lock = threading.Lock()
_query_model = SentenceTransformer("all-MiniLM-L6-v2")


def _on_doc_change(key, row, time, is_addition):
    if not is_addition:
        with _live_lock:
            for chunk_key in [k for k in _live_chunks if k.startswith(f"{key}::")]:
                _live_chunks.pop(chunk_key, None)
            _live_chunks.pop(str(key), None)
            _rag_state["chunks_indexed"] = len(_live_chunks)
        return

    # The DAG already extracted text per format, so this is always UTF-8.
    if isinstance(row, dict):
        raw = row.get("data", row.get("text", ""))
        metadata = row.get("_metadata", {})
    else:
        raw = getattr(row, "data", "")
        metadata = getattr(row, "_metadata", {})

    text = raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else (raw or "")
    if not text or len(text.strip()) <= 10:
        return

    path = _metadata_path(metadata)
    meta = {"path": path, "source": "stream"} if path else {"source": "stream"}

    # A whole document is far larger than the embedder's window, so index it in
    # the same word-based chunks the preloader uses.
    words = text.split()
    chunks = [
        " ".join(words[i:i + 250])
        for i in range(0, len(words), 250)
    ]
    chunks = [c for c in chunks if len(c) > 50]
    if not chunks:
        return

    embeddings = _query_model.encode(chunks, convert_to_numpy=True)

    with _live_lock:
        for stale in [k for k in _live_chunks if k.startswith(f"{key}::")]:
            _live_chunks.pop(stale, None)
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            _live_chunks[f"{key}::{i}"] = {
                "text": chunk[:800],
                "metadata": meta,
                "embedding": emb,
            }
        _rag_state["chunks_indexed"] = len(_live_chunks)
        _rag_state["store_status"] = "active"
        _rag_state["last_reindex"] = datetime.now(timezone.utc).strftime("%H:%M:%S")


pw.io.subscribe(_policy_docs, on_change=_on_doc_change)


# --- Preload txt files directly so retrieval works immediately ---

def _preload_policies():
    """Index every supported policy document up front.

    Previously only .txt was preloaded, so a PDF policy contributed nothing to
    retrieval until (and unless) the stream picked it up.
    """
    for f in sorted(os.listdir(POLICY_DIR)):
        p = os.path.join(POLICY_DIR, f)
        if not os.path.isfile(p):
            continue
        if os.path.splitext(f)[1].lower() not in SUPPORTED_EXTENSIONS:
            continue

        try:
            with open(p, "rb") as fp:
                data = fp.read()

            text = extract_document_text(p, data)
            if not text.strip():
                continue

            words = text.split()
            chunks = [" ".join(words[i:i + 250]) for i in range(0, len(words), 250)]
            chunks = [c for c in chunks if len(c) > 50]
            if not chunks:
                continue

            embeddings = _query_model.encode(chunks, convert_to_numpy=True)
            with _live_lock:
                for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
                    _live_chunks[f"preload_{f}_{i}"] = {
                        "text": chunk[:800],
                        "metadata": {"path": p, "source": "preload"},
                        "embedding": emb,
                    }
            print(f"[RAG] preloaded {len(chunks)} chunks from {f}")
        except Exception as e:
            _record_parse_error(f, f"{type(e).__name__}: {str(e)[:160]}")

    with _live_lock:
        _rag_state["chunks_indexed"] = len(_live_chunks)
        if _live_chunks:
            _rag_state["store_status"] = "active"
            _rag_state["last_reindex"] = datetime.now(timezone.utc).strftime("%H:%M:%S")
            print(f"[RAG] preloaded {len(_live_chunks)} chunks")

threading.Thread(target=_preload_policies, daemon=True).start()


# --- Retrieval ---

def retrieve_policy_context(query, k=2):
    _scan_policy_files()

    with _live_lock:
        if not _live_chunks:
            return {
                "context": "Policy index initializing...",
                "policy_file": "loading...",
                "similarity_score": 0.0,
                "index_type": "Pathway DocumentStore (Initializing)",
                "policy_last_updated": _sync_age(),
                "docs_indexed": _rag_state["docs_indexed"],
                "embed_model": "all-MiniLM-L6-v2",
            }

        q_emb = _query_model.encode([query], convert_to_numpy=True)

        keys = list(_live_chunks.keys())
        embeddings = np.array([_live_chunks[k_]["embedding"] for k_ in keys])
        scores = (embeddings @ q_emb.T).flatten()
        topk_idx = np.argsort(scores)[-k:][::-1]

        best = _live_chunks[keys[topk_idx[0]]]
        ctx = "\n\n".join([_live_chunks[keys[i]]["text"][:400] for i in topk_idx])

        meta = best.get("metadata", {})
        fname = os.path.basename(meta.get("path", "policy-document")) if isinstance(meta, dict) else "policy-document"

        return {
            "context": ctx[:800],
            "policy_file": fname,
            "similarity_score": round(float(scores[topk_idx[0]]), 4),
            "index_type": "Pathway DocumentStore (Live Hybrid Index)",
            "policy_last_updated": _sync_age(),
            "docs_indexed": _rag_state["docs_indexed"],
            "embed_model": "all-MiniLM-L6-v2",
        }


def _sync_age():
    lr = _rag_state.get("last_reindex")
    return f"Last index: {lr} UTC" if lr else "Initializing..."


def get_governance_rule():
    return (
        f"AQI >= {HIGH_AQI_THRESHOLD} | "
        f"{PERSISTENCE_THRESHOLD} Consecutive Windows | "
        f"3min Sliding | 1min Hop | "
        f"Hysteresis: 2 confirmations | "
        f"Protocol: CAQM GRAP Escalation"
    )


# --- Advisory generation ---

def generate_grounded_advisory(
    aqi, level, grap_description, band, fire_count,
    high_count=0, remaining_windows=0, projected_time="N/A",
    transport_score=0, transport_label="none",
    wind_speed=None, wind_dir=None,
):
    rag = retrieve_policy_context(f"{level} {band} GRAP enforcement CPCB")
    rule = get_governance_rule()

    legal = (
        f"LEGAL BASIS\n{'='*50}\n"
        f"CPCB Band  : {band}\n"
        f"GRAP Stage : {level}\n"
        f"Action     : {grap_description}\n"
    )

    signal = (
        f"\nLIVE SIGNAL\n{'='*50}\n"
        f"AQI              : {aqi}\n"
        f"Persistence      : {high_count} windows (Threshold: {PERSISTENCE_THRESHOLD})\n"
        f"Remaining        : {remaining_windows}\n"
        f"Projected Trigger: {projected_time}\n"
        f"Fire Hotspots    : {fire_count}\n"
    )

    gov = f"\nTRIGGER RULE\n{'='*50}\n{rule}\n"

    if high_count >= PERSISTENCE_THRESHOLD:
        esc = (
            f"\nESCALATION: TRIGGERED\n{'='*50}\n"
            f"{high_count} consecutive windows >= {HIGH_AQI_THRESHOLD}.\n"
            f"Immediate regulatory activation required.\n"
        )
        enf = (
            f"\nMANDATORY ACTIONS\n{'='*50}\n"
            f"- Construction/demolition restrictions\n"
            f"- High-emission vehicle entry ban\n"
            f"- Industrial compliance verification\n"
            f"- Public health advisory issuance\n"
            f"- School outdoor activity suspension\n"
        )
    else:
        esc = (
            f"\nESCALATION: WATCH\n{'='*50}\n"
            f"Threshold not met. {remaining_windows} windows remaining.\n"
            f"Projected trigger: {projected_time}\n"
        )
        enf = (
            f"\nPREPARED PROTOCOL\n{'='*50}\n"
            f"- Construction restriction readiness\n"
            f"- Vehicle enforcement standby\n"
            f"- Public health advisory drafted\n"
        )

    if transport_label == "regional_transport":
        ws = f"{wind_speed:.1f}" if wind_speed else "N/A"
        wd = f"{wind_dir:.0f}" if wind_dir else "N/A"
        causal = (
            f"\nCAUSAL ATTRIBUTION\n{'='*50}\n"
            f"Satellite-detected thermal anomalies upwind.\n"
            f"Transport Score  : {transport_score}/100\n"
            f"Wind             : {ws} m/s from {wd} deg\n"
            f"Source           : NASA FIRMS VIIRS_SNPP_NRT\n"
        )
    elif transport_label == "possible_transport":
        causal = (
            f"\nCAUSAL ATTRIBUTION\n{'='*50}\n"
            f"Limited upwind thermal activity detected.\n"
            f"Transport Score  : {transport_score}/100\n"
        )
    else:
        causal = (
            f"\nCAUSAL ATTRIBUTION\n{'='*50}\n"
            f"No upwind thermal anomalies. Local emission dominant.\n"
        )

    pol = (
        f"\nPOLICY SOURCE ({rag['index_type']})\n{'='*50}\n"
        f"Document  : {rag['policy_file']}\n"
        f"Score     : {rag['similarity_score']}\n"
        f"Sync      : {rag['policy_last_updated']}\n"
        f"Indexed   : {rag['docs_indexed']} documents\n"
        f"Chunks    : {_rag_state['chunks_indexed']}\n"
        f"Embedder  : {rag['embed_model']}\n"
    )

    return {
        "advisory": legal + signal + gov + esc + enf + causal + pol,
        "policy_file": rag["policy_file"],
        "similarity_score": rag["similarity_score"],
        "policy_last_updated": rag["policy_last_updated"],
        "index_type": rag["index_type"],
        "docs_indexed": rag["docs_indexed"],
        "embed_model": rag["embed_model"],
        "governance_rule": rule,
    }
